import json
from pathlib import Path

from app.services.education_domain import (
    classify_combined_heading,
    classify_section_heading,
    extract_education_entities,
    segment_resume,
    validate_education_record,
    validate_experience_record,
)
from app.services.nlp_pipeline import extract_education, extract_experience_records
from app.services.nlp_pipeline import extract_contact_info, extract_certifications
from app.services.extractor import _clean_extracted_text, extract_text_from_docx
from app.services.recommender import extract_resume_skills


FIXTURE_PATH = Path(__file__).parent / "fixtures" / "education_resume_cases.json"


def test_section_classifier_handles_education_domain_and_combined_headings():
    assert classify_section_heading("TEACHING EXPERIENCE") == "experience"
    assert classify_section_heading("EDUCATIONAL BACKGROUND") == "education"
    assert classify_section_heading("COMPÉTENCES TECHNIQUES") == "skills"
    assert classify_combined_heading("CERTIFICATIONS, SKILLS & AWARDS") == [
        "certifications", "skills", "awards"
    ]


def test_segmenter_preserves_ordered_section_context():
    sections = segment_resume("MARIA\nSKILLS\nLesson Planning\nEDUCATION\nBachelor of Education")
    assert [section["section"] for section in sections] == ["header", "skills", "education"]
    assert sections[1]["lines"] == ["Lesson Planning"]


def test_hybrid_education_entities_include_evidence_and_confidence():
    entities = extract_education_entities(
        "TEACHING EXPERIENCE\nGrade 8 English Teacher\nUsed differentiated instruction.\n"
        "CERTIFICATIONS\nLicensed Professional Teacher"
    )
    entity_types = {entity["type"] for entity in entities}
    assert {"education_role", "grade_level", "subject", "education_skill", "license"}.issubset(entity_types)
    assert all(entity["evidence"] and entity["confidence"] in {"high", "medium"} for entity in entities)


def test_structured_validators_reject_noise_but_allow_unknown_employer():
    assert validate_experience_record({
        "job_title": "English Teacher", "company": "Not Identified", "years": 2
    })[0]
    assert not validate_experience_record({
        "job_title": "Prepared lessons for Grade 8 students",
        "company": "(A.Y", "years": 1,
    })[0]
    assert validate_education_record({
        "degree": "Bachelor of Secondary Education",
        "institution": "Example University",
    })[0]
    assert not validate_education_record({
        "degree": "High School With High Honors",
        "institution": "Unknown Institution",
    })[0]


def test_regression_resume_corpus():
    cases = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    for case in cases:
        skills = extract_resume_skills(case["text"])
        experience = extract_experience_records(case["text"])
        education = extract_education(case["text"])

        for expected in case["expected_skills"]:
            assert expected in skills, (case["id"], expected, skills)
        assert [record["job_title"] for record in experience] == case["expected_roles"], case["id"]
        assert len(education) == case["expected_education_count"], case["id"]


def test_decorative_template_headings_preserve_section_boundaries_and_identity():
    raw = """R O D E L B . P U N Z A L A N , L P T
E D U C A T I O N
Elementary Graduate
Sto. Rosario Elementary School (2006-2012)
W O R K E X P E R I E N C E
Colegio De Sebastian-Pampanga Inc. (2022-2025)
VOLLEYBALL COACH AND TRAINER
P U B L I C A T I O N S & P R E S E N T A T I O N S
Teacher Education Student Conference
November 25, 2021
Cebu Normal University
Rodel B. Punzalan, LPT
APPLICANT'S SIGNATURE
"""
    text = _clean_extracted_text(raw)

    assert "WORK EXPERIENCE" in text
    assert "PUBLICATIONS & PRESENTATIONS" in text
    assert extract_contact_info(text)["name"] == "Rodel B. Punzalan"
    records = extract_experience_records(text)
    assert {record["job_title"] for record in records} == {"VOLLEYBALL COACH AND TRAINER"}
    assert all("Conference" not in record["job_title"] for record in records)


def test_combined_license_heading_extracts_wrapped_dated_credentials():
    text = """LICENSES & CERTIFICATIONS
1Seminar Workshop on Senior High School on PPST and UDL
Training
July 19, 2024
DEPED Region 3 Coaches Accreditation
October 5, 2024
CHARACTER REFERENCES
Maria Santos
"""
    records = extract_certifications(text)
    assert records == [
        {
            "certification_name": "Seminar Workshop on Senior High School on PPST and UDL Training",
            "credential_type": "Certification",
            "issuer": None,
            "date_obtained": "July 19, 2024",
        },
        {
            "certification_name": "DEPED Region 3 Coaches Accreditation",
            "credential_type": "Certification",
            "issuer": None,
            "date_obtained": "October 5, 2024",
        },
    ]


def test_template_education_levels_and_wrapped_skills_are_not_cross_assigned():
    text = """EDUCATION
Elementary Graduate
Sto. Rosario Elementary School (2006-2012)
High School Graduate
Mexico National High School
Junior High School (2012-2016)
Senior High School (2016-2018)
College Graduate
Don Honorio Ventura State University (2018-2022)
Graduate School
Pampanga State University (Completed Academic Requirement)
SKILLS
Teamwork
Verbal & Written
communication
CONTACT
https://www.researchgate.net/profile/example
"""
    education = extract_education(text)
    skills = extract_resume_skills(text)
    assert {record["institution"] for record in education} >= {
        "Sto. Rosario Elementary School", "Mexico National High School",
        "Don Honorio Ventura State University", "Pampanga State University",
    }
    assert [record["degree"] for record in education] == [
        "Elementary Graduate", "High School Graduate", "College Graduate",
        "Graduate School (Completed Academic Requirement)",
    ]
    assert "Verbal & Written communication" in skills
    assert "net" not in skills


def test_late_column_name_wrapped_degree_and_role_first_jobs():
    text = """Committed teacher helping students achieve their fullest potential.
WORK EXPERIENCE
Public Senior High School Teacher
August 29, 2023 - present
Pampanga High School
San Fernando, Pampanga, Philippines
Senior High School Teacher
January 13, 2021 - May 31, 2023 (Blended Learning Modality)
Colegio De Sebastian Pampanga Inc.
JOYCE ANN R. PINEDA
CONTACT
joyceannpineda14@gmail.com
0955-049-0222
EDUCATION
Bachelor of Secondary
Education Major in Biological
Science
- City College of San Fernando
SKILLS
- Classroom management
- Lesson planning and curriculum development
"""
    assert extract_contact_info(text)["name"] == "Joyce Ann R. Pineda"
    assert extract_education(text) == [{
        "degree": "Bachelor of Secondary Education Major in Biological Science",
        "institution": "City College of San Fernando",
    }]
    assert [(row["job_title"], row["company"]) for row in extract_experience_records(text)] == [
        ("Public Senior High School Teacher", "Pampanga High School"),
        ("Senior High School Teacher", "Colegio De Sebastian Pampanga Inc"),
    ]
    assert extract_resume_skills(text) == [
        "Classroom management", "Lesson planning and curriculum development"
    ]


def test_work_related_heading_stops_achievements_and_mobile_format_is_supported():
    text = """JAYSON E. TUNGOL
Mobile: 0997-7301-789
WORK-RELATED EXPERIENCE
Senior High School Math Teacher
San Isidro High School
June 15, 2016 up to present
ACHIEVEMENTS
SBM Coordinator
San Isidro High School 2014 - 2016
"""
    assert extract_contact_info(text)["phone"] == "0997-7301-789"
    assert [(row["job_title"], row["company"]) for row in extract_experience_records(text)] == [
        ("Senior High School Math Teacher", "San Isidro High School")
    ]


def test_application_form_docx_tables_are_extracted_by_schema(tmp_path):
    from docx import Document

    document = Document()
    personal = document.add_table(rows=2, cols=4)
    personal.rows[0].cells[0].text = "Full Name"
    personal.rows[0].cells[1].text = "MYCHEL B. PELAYO"
    personal.rows[0].cells[2].text = "Residence"
    personal.rows[0].cells[3].text = "Pampanga"
    personal.rows[1].cells[0].text = "Cellphone No."
    personal.rows[1].cells[1].text = "0925 664 2375"
    personal.rows[1].cells[2].text = "Email Address"
    personal.rows[1].cells[3].text = "mychel@example.com"

    education = document.add_table(rows=3, cols=5)
    for cell, value in zip(education.rows[0].cells, ["Level", "School Name", "Kurso", "Location", "Year Graduated"]):
        cell.text = value
    for cell, value in zip(education.rows[1].cells, ["Elementary", "San Isidro Elem. School", "N/A", "Bacolor", "1982"]):
        cell.text = value
    for cell, value in zip(education.rows[2].cells, ["College", "DHVCAT", "BSIE", "Bacolor", "1990"]):
        cell.text = value

    experience = document.add_table(rows=2, cols=6)
    for cell, value in zip(experience.rows[0].cells, ["Company / Organization", "Position", "Location", "From", "To", "Primary Duties"]):
        cell.text = value
    for cell, value in zip(experience.rows[1].cells, ["DepED", "Teacher V", "Bacolor", "1995", "present", "Teaching Math"]):
        cell.text = value

    path = tmp_path / "form_resume.docx"
    document.save(path)
    text = extract_text_from_docx(path)
    assert extract_contact_info(text)["name"] == "Mychel B. Pelayo"
    assert extract_education(text) == [
        {"degree": "Elementary Education", "institution": "San Isidro Elem. School"},
        {"degree": "BSIE", "institution": "DHVCAT"},
    ]
    assert extract_experience_records(text)[0]["company"] == "DepED"

