"""
Unit tests for the Resume Screening System core logic.

Coverage:
  - recommender.analyze_skills           (skill matching, aliases, edge cases)
  - recommender.analyze_preferred_skills (bonus capping, empty lists)
  - recommender.generate_recommendation  (all threshold branches)
  - recommender.get_degree_rank          (all degree levels + unknowns)
  - recommender.generate_analysis_narrative (all tier / disqualified branches)
  - recommender.evaluate_candidate       (integration, critical skill enforcement)
  - matching_engine.calculate_fit_score  (weight math, capping)
  - matching_engine.calculate_text_similarity (basic similarity behaviour)
  - matching_engine.calculate_skills_match    (percentage match helper)
  - utils.files.unique_upload_filename   (uniqueness, safety, extension)
  - utils.files.is_path_inside_directory (containment check)
  - utils.files.safe_delete_uploaded_file (delete guard)
"""

import os
import tempfile
import unittest
from datetime import datetime
from types import SimpleNamespace

from app.services.evidence import build_candidate_evidence
from app.services.matching_engine import (
    FIT_WEIGHT_PERCENTS,
    calculate_fit_score,
    calculate_skills_match,
    calculate_text_similarity,
)
from app.services.extractor import (
    extract_text_from_docx,
    extract_text_from_pdf,
    _extract_page_columns,
    _clean_extracted_text,
    _select_best_pdf_text,
    _text_quality_score,
)
from app.services.nlp_pipeline import (
    extract_contact_info,
    extract_certifications,
    extract_education,
    extract_experience_records,
    extract_years_of_experience,
)
from app.services.recommender import (
    SKILL_ALIASES,
    analyze_preferred_skills,
    analyze_skills,
    extract_resume_skills,
    estimate_decision_confidence,
    evaluate_candidate,
    generate_analysis_narrative,
    generate_decision_explanation,
    generate_recommendation,
    get_degree_rank,
)
from app.utils.files import (
    job_upload_directory,
    is_path_inside_directory,
    safe_delete_uploaded_file,
    unique_upload_filename,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_narrative(**kwargs):
    """Call generate_analysis_narrative with safe defaults so individual tests
    only need to specify the fields they care about."""
    defaults = dict(
        job_title="Software Engineer",
        fit_score=80.0,
        skill_score=90.0,
        exp_score=80.0,
        edu_score=100.0,
        matched_skills=["Python", "SQL"],
        missing_skills=[],
        total_exp_years=3.0,
        extracted_edu=[{"degree": "Bachelor's", "institution": "State U"}],
        experience_req=2,
        education_req="Bachelor's",
        disqualified_by_critical_skills=False,
        matched_preferred=None,
        preferred_bonus=0.0,
    )
    defaults.update(kwargs)
    return generate_analysis_narrative(**defaults)


# ===========================================================================
# analyze_skills
# ===========================================================================

class TestExtractExperienceRecords(unittest.TestCase):

    def test_skill_bullets_without_dates_are_not_misread_as_job_entries(self):
        sample = """
Work Experience
- Strong hands-on experience in developing enterprise applications using Java/J2EE
- Service Oriented Architecture (SOA), REST APIs, SQL Developer
- IBM Rational Application Developer, Eclipse, IntelliJ IDEA
Excellent team player with strong leadership abilities
"""
        records = extract_experience_records(sample)
        self.assertEqual(records, [])

    def test_month_year_present_range_is_recognized(self):
        sample = """
Work Experience
CVS, Woonsocket, Rhode Island                                 Full Stack Java Developer
April 2016 - Present
Responsibilities:
Built web apps.

Toll Brothers, Horsham Township, Pennsylvania                  Software Engineer
December 2015 - March 2016
Responsibilities:
Worked on backend services.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]['company'], 'CVS')
        self.assertIn('Developer', records[0]['job_title'])

    def test_client_location_role_resume_format_is_recognized(self):
        sample = """
PROFESSIONAL EXPERIENCE
Client: Capital One                                      Mar 17 - Till date
Location: Mclean, VA
Role:  Full Stack Java Developer
Responsibilities:
Built banking applications.

Client: Cigna Healthcare                                 Nov 13- Dec 15 Location: Windsor, CT
Role: Java Developer
Responsibilities:
Developed web services.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]['company'], 'Capital One')
        self.assertEqual(records[0]['location'], 'Mclean, VA')
        self.assertEqual(records[0]['job_title'], 'Full Stack Java Developer')
        self.assertEqual(records[1]['company'], 'Cigna Healthcare')
        self.assertEqual(records[1]['location'], 'Windsor, CT')

    def test_explicit_total_years_preferred_over_till_date_sum(self):
        sample = """
OBJECTIVE:
Around 8 years of strong software experience.
PROFESSIONAL EXPERIENCE
Client: Capital One                                      Mar 17 - Till date
Location: Mclean, VA
Role: Full Stack Java Developer
"""
        self.assertEqual(extract_years_of_experience(sample), 8.0)

    def test_email_lines_are_not_extracted_as_companies(self):
        sample = """
Jane Santos
jane.santos@gmail.com

Professional Experience
Software Engineer
jane.santos@gmail.com
January 2020 - January 2022
Responsibilities:
Built internal applications.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertNotIn("gmail", records[0]["company"].lower())

    def test_education_school_designation_format_is_recognized(self):
        sample = """
Teaching Experience
School: San Isidro National High School                  June 2019 - Present
Location: Quezon City
Designation: Mathematics Teacher
Responsibilities:
Prepared lesson plans and assessed student performance.

Institution: Bright Minds Learning Center                May 2017 - March 2019
Location: Manila
Position: Tutor
Responsibilities:
Provided after-school academic support.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["company"], "San Isidro National High School")
        self.assertEqual(records[0]["location"], "Quezon City")
        self.assertEqual(records[0]["job_title"], "Mathematics Teacher")
        self.assertEqual(records[1]["company"], "Bright Minds Learning Center")
        self.assertEqual(records[1]["job_title"], "Tutor")

    def test_inline_education_role_format_is_recognized(self):
        sample = """
Professional Experience
University of Makati                                  Guidance Counselor
June 2020 - Present
Responsibilities:
Supported students with academic and personal concerns.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["company"], "University of Makati")
        self.assertEqual(records[0]["job_title"], "Guidance Counselor")

    def test_short_non_tech_role_with_numeric_dates_is_preserved(self):
        sample = """
Work Experience
Recruitment Assistant
PeopleFirst Inc.
01/2024 - 06/2024
Assisted with applicant records and interviews.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "Recruitment Assistant")
        self.assertEqual(records[0]["company"], "PeopleFirst Inc")
        self.assertGreater(records[0]["years"], 0)
        self.assertLess(records[0]["years"], 1)

    def test_pipe_separated_retail_role_company_and_location_are_recognized(self):
        sample = """
Employment History
Cashier | FreshMart Retail | Makati
June 2022 - August 2023
Handled customer payments and daily cash reports.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "Cashier")
        self.assertEqual(records[0]["company"], "FreshMart Retail")
        self.assertEqual(records[0]["location"], "Makati")

    def test_hyphen_separated_title_company_and_dates_are_recognized(self):
        sample = """
Work Experience
Software Engineer - Acme Digital Solutions - January 2020 - Present
Built APIs and internal reporting tools.

Project Manager - Northwind Services - March 2018 - December 2019
Led delivery planning and client coordination.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["job_title"], "Software Engineer")
        self.assertEqual(records[0]["company"], "Acme Digital Solutions")
        self.assertEqual(records[1]["job_title"], "Project Manager")
        self.assertEqual(records[1]["company"], "Northwind Services")

    def test_company_title_date_inline_format_is_recognized(self):
        sample = """
Professional Experience
Acme Digital Solutions - Software Engineer - January 2020 - Present
Northwind Services, Project Manager, March 2018 - December 2019
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["job_title"], "Software Engineer")
        self.assertEqual(records[0]["company"], "Acme Digital Solutions")
        self.assertEqual(records[1]["job_title"], "Project Manager")
        self.assertEqual(records[1]["company"], "Northwind Services")

    def test_parenthesized_dates_do_not_pollute_job_title(self):
        sample = """
Experience
Software Engineer, Acme Digital Solutions (January 2020 - Present)
Built APIs and dashboards.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "Software Engineer")
        self.assertEqual(records[0]["company"], "Acme Digital Solutions")

    def test_pdf_split_year_is_normalized_for_experience_dates(self):
        sample = _clean_extracted_text("""
Professional Experience
Bank of Utah, Ogden, Utah                                                           J2EE Developer
May  2013  - October 201 4
Responsibilities:
Developed Java applications.
""")
        records = extract_experience_records(sample)

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "J2EE Developer")
        self.assertEqual(records[0]["company"], "Bank of Utah")
        self.assertEqual(records[0]["location"], "Ogden, Utah")

    def test_pdf_state_split_into_title_is_restored_to_location(self):
        sample = _clean_extracted_text("""
Professional Experience
Toll Brothers, Horsham Township,  Pennsylvania                  Software Engineer
December 2015 -  March 2016
Responsibilities:
Developed JSP applications.
""")
        records = extract_experience_records(sample)

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "Software Engineer")
        self.assertEqual(records[0]["company"], "Toll Brothers")
        self.assertEqual(records[0]["location"], "Horsham Township, Pennsylvania")

    def test_education_dates_are_not_counted_as_work_experience(self):
        sample = """
Education
Bachelor of Science in Accounting
State University
2018 - 2022
Four years of full-time study.
"""
        self.assertEqual(extract_years_of_experience(sample), 0.0)

    def test_inline_company_title_never_uses_end_date_as_title(self):
        sample = """
Professional Experience
Toll Brothers, Horsham Township, Pennsylvania                  Software Engineer
December 2015 - March 2016
Responsibilities:
Developed internal applications.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["job_title"], "Software Engineer")
        self.assertEqual(records[0]["company"], "Toll Brothers")
        self.assertEqual(records[0]["location"], "Horsham Township, Pennsylvania")

    def test_company_date_then_role_format_creates_separate_jobs(self):
        sample = """
Professional Experience
Wind Stream Communication, Dallas, TX                 SEP-2016 - TILL DATE
Role: Sr Java Programmer
Responsibilities:
Built web services.
IBM DALLAS,TEXASSEP 2015-AUG 2016
Role: Sr Java Programmer
Responsibilities:
Maintained enterprise applications.
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["company"], "Wind Stream Communication")
        self.assertEqual(records[0]["location"], "Dallas, TX")
        self.assertEqual(records[1]["company"], "IBM")
        self.assertEqual(records[1]["location"], "DALLAS, TEXAS")

    def test_consecutive_client_role_blocks_are_not_merged(self):
        sample = """
Professional Experience
Client: Datacard Software India Pvt Ltd, Bangalore    Jan 2003 - July 2007
Role: Java Developer, Performance Engineer, Lead in Java based
Product Title: Affina
Client: Infosys Technologies Pvt Ltd, Bangalore       August 2002 to January 2003
Role: Software Automation Test Engineer
"""
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["company"], "Datacard Software India Pvt Ltd")
        self.assertIn("Java Developer", records[0]["job_title"])
        self.assertEqual(records[1]["company"], "Infosys Technologies Pvt Ltd")

    def test_complete_history_is_not_truncated_to_five_roles(self):
        entries = []
        company_names = ["Alpha", "Bravo", "Charlie", "Delta", "Echo", "Foxtrot"]
        for offset, company_name in enumerate(company_names):
            year = 2010 + offset
            entries.append(
                f"{company_name} Company  Software Engineer\nJanuary {year} - January {year + 1}"
            )
        sample = "Professional Experience\n" + "\n\n".join(entries)
        records = extract_experience_records(sample)
        self.assertEqual(len(records), 6)

    def test_explicit_it_experience_statement_is_preferred(self):
        sample = """
Professional Summary
17 years of IT experience in software delivery and program management.
Professional Experience
Current Company  Program Manager
May, 2017 to Till Date
"""
        self.assertEqual(extract_years_of_experience(sample), 17.0)


class TestExtractContactInfo(unittest.TestCase):

    def test_anonymized_teacher_candidate_does_not_use_skill_as_name(self):
        sample = """
Teacher Candidate 2
Summary
Enthusiastic teacher effective at fostering a positive learning environment.
Skills
Critical thinker
Calm under pressure
Education
Masters: Teaching
"""
        contact = extract_contact_info(sample)

        self.assertEqual(contact["name"], "Unknown Candidate")

    def test_document_format_word_is_not_used_as_candidate_name(self):
        sample = """
Word
word@example.com

Professional Summary
Experienced applicant with strong communication and organizational skills.
"""
        contact = extract_contact_info(sample)

        self.assertEqual(contact["name"], "Unknown Candidate")

    def test_name_split_across_pdf_header_lines_beats_narrative_phrase(self):
        sample = """
LETICIA
AKILAL
Recherche d'un contrat d'apprentissage en IT dans le cadre
d'une L3 MIAGE
E-mail : akilallaeticia@gmail.com
"""
        contact = extract_contact_info(sample)

        self.assertEqual(contact["name"], "Leticia Akilal")

    def test_doubled_pdf_glyphs_are_collapsed_in_candidate_name(self):
        sample = """
BBOORRHHAANN GGHHEENNNNAAII
TTEECCHHNNIICCIIEENN SSUUPPÉÉRRIIEEUURR
borhan.g@live.fr
"""
        contact = extract_contact_info(sample)

        self.assertEqual(contact["name"], "Borhan Ghennai")

    def test_pdf_split_first_name_is_merged_using_email_identity(self):
        sample = """
Name: Abi ral Pandey
Email: abiral.pandey88@gmail.com
Phone: 940-242-3303
"""
        contact = extract_contact_info(sample)

        self.assertEqual(contact["name"], "Abiral Pandey")

    def test_tech_word_is_not_used_as_candidate_name(self):
        sample = """
Java Developer
juan.delacruz@gmail.com
09171234567

Professional Summary
Experienced Java developer with Spring Boot exposure.
"""
        contact = extract_contact_info(sample)
        self.assertNotEqual(contact["name"], "Java")
        self.assertNotEqual(contact["name"], "Java Developer")
        self.assertEqual(contact["name"], "Juan Delacruz")

    def test_header_name_is_preferred_over_role_and_email(self):
        sample = """
Maria Santos
Full Stack Java Developer
maria.santos@gmail.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Maria Santos")

    def test_teacher_role_is_not_used_as_candidate_name(self):
        sample = """
Teacher
ana.reyes@gmail.com
09181234567

Professional Summary
Licensed educator with classroom management experience.
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Ana Reyes")

    def test_location_label_is_not_used_as_candidate_name(self):
        sample = """
Teacher
ana.reyes@gmail.com
09181234567

Teaching Experience
School: San Isidro National High School
Location: Quezon City
Designation: Mathematics Teacher
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Ana Reyes")

    def test_all_caps_name_and_philippine_phone_are_normalized(self):
        sample = """
MARIA DELA CRUZ
Licensed Professional Teacher
maria.delacruz@example.com
+63 917 123 4567
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Maria Dela Cruz")
        self.assertEqual(contact["phone"], "+63 917 123 4567")

    def test_lowercase_explicit_name_is_accepted(self):
        sample = """
Full Name: juan miguel santos
Email: juan.santos@example.com
Mobile: 0918 123 4567
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Juan Miguel Santos")
        self.assertEqual(contact["phone"], "0918 123 4567")

    def test_company_and_location_header_fall_back_to_email_name(self):
        sample = """
QUEZON CITY
ABC CORPORATION
Software Engineer
paolo.reyes@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Paolo Reyes")

    def test_single_name_is_accepted_when_email_confirms_it(self):
        sample = """
Achyuth
540-999-8048
achyuth.java88@gmail.com
OBJECTIVE:
Around 8 years of software experience.
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Achyuth")

    def test_single_name_is_supported_by_concatenated_email_identity(self):
        sample = """
Anudeep
Sr Java Programmer
anudeepreddynallamada@gmail.com
Professional Summary:
Eight years of software development experience.
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Anudeep")

    def test_last_name_first_explicit_label_is_reordered(self):
        sample = """
Applicant Name: SANTOS, JUAN DELA CRUZ
Email: juan.santos@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Juan Dela Cruz Santos")

    def test_honorific_credential_apostrophe_and_hyphen_are_formatted(self):
        sample = """
Full Name: Dr. ANNE-MARIE O'NEIL, RN
Email: anne.oneil@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Anne-Marie O'Neil")

    def test_name_is_extracted_from_mixed_header_line(self):
        sample = """
ALYSSA MAE RAMOS | Licensed Professional Teacher | alyssa.ramos@example.com
0917 555 0101
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Alyssa Mae Ramos")

    def test_curriculum_vitae_of_name_is_recognized(self):
        sample = """
CURRICULUM VITAE OF LUZVIMINDA REYES
Teacher
luzviminda.reyes@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Luzviminda Reyes")

    def test_concatenated_email_supports_real_name_over_location(self):
        sample = """
Makati Philippines
JOHN MICHAEL SANTOS
Software Engineer
johnmichaelsantos@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "John Michael Santos")

    def test_non_tech_role_and_company_are_not_candidate_name(self):
        sample = """
Senior Accountant
Bright Future Solutions
roberto.garcia@example.com
"""
        contact = extract_contact_info(sample)
        self.assertEqual(contact["name"], "Roberto Garcia")


class TestCandidateEvidence(unittest.TestCase):

    def test_evidence_uses_resume_sections_and_requirement_status(self):
        sample = """
Professional Summary
Software engineer with 4 years of professional experience.
Technical Skills
Python, JS, SQL, Docker
Professional Experience
Software Engineer | Acme Corporation | Manila
January 2020 - January 2024
Education
Bachelor of Science in Computer Science
State University
"""
        experience = [SimpleNamespace(
            job_title="Software Engineer",
            company="Acme Corporation",
            location="Manila",
            years=4.0,
        )]
        education = [SimpleNamespace(
            degree="Bachelor of Science in Computer Science",
            institution="State University",
        )]

        evidence = build_candidate_evidence(
            resume_text=sample,
            matched_skills=["Python", "JavaScript"],
            missing_skills=["Kubernetes"],
            matched_critical_skills=["SQL"],
            missing_critical_skills=["AWS"],
            matched_preferred_skills=["Docker"],
            experience_records=experience,
            education_records=education,
            experience_requirement=3,
            education_requirement="Bachelor's",
        )

        javascript = next(item for item in evidence["skills"] if item["name"] == "JavaScript")
        self.assertEqual(javascript["section"], "Skills")
        self.assertIn("JS", javascript["excerpt"])
        self.assertIn("Acme Corporation", evidence["experience"][0]["excerpt"])
        self.assertIn("State University", evidence["education"][0]["excerpt"])
        self.assertTrue(evidence["experience_status"]["met"])
        self.assertTrue(evidence["education_status"]["met"])
        self.assertEqual(
            {(item["kind"], item["name"]) for item in evidence["missing"]},
            {("Critical", "AWS"), ("Required", "Kubernetes")},
        )

    def test_evidence_marks_unmet_requirements(self):
        evidence = build_candidate_evidence(
            resume_text="Professional Summary\nOne year of professional experience.",
            matched_skills=[],
            missing_skills=["Excel"],
            matched_critical_skills=[],
            missing_critical_skills=[],
            matched_preferred_skills=[],
            experience_records=[],
            education_records=[],
            experience_requirement=3,
            education_requirement="Bachelor's",
        )
        self.assertFalse(evidence["experience_status"]["met"])
        self.assertFalse(evidence["education_status"]["met"])

    def test_work_evidence_excludes_duties_and_next_section_text(self):
        sample = """
Work Experience
Software Engineer
Acme Corporation
January 2021 - January 2024
Responsibilities: Managed payroll records and prepared weekly reports
Education
Bachelor of Science in Computer Science
State University
"""
        experience = [SimpleNamespace(
            job_title="Software Engineer", company="Acme Corporation",
            location="Manila", years=3.0,
        )]
        evidence = build_candidate_evidence(
            resume_text=sample, matched_skills=[], missing_skills=[],
            matched_critical_skills=[], missing_critical_skills=[],
            matched_preferred_skills=[], experience_records=experience,
            education_records=[],
        )

        excerpt = evidence["experience"][0]["excerpt"]
        self.assertIn("Software Engineer", excerpt)
        self.assertIn("Acme Corporation", excerpt)
        self.assertIn("January 2021 - January 2024", excerpt)
        self.assertNotIn("payroll", excerpt.lower())
        self.assertNotIn("Bachelor", excerpt)

    def test_education_evidence_excludes_honors_and_work_text(self):
        sample = """
Education
Bachelor of Science in Information Technology
State University
2018 - 2022
Awards: Dean's Lister and Best Capstone Project
Work Experience
Technical Support Specialist
Northwind Services
"""
        education = [SimpleNamespace(
            degree="Bachelor of Science in Information Technology",
            institution="State University",
        )]
        evidence = build_candidate_evidence(
            resume_text=sample, matched_skills=[], missing_skills=[],
            matched_critical_skills=[], missing_critical_skills=[],
            matched_preferred_skills=[], experience_records=[],
            education_records=education,
        )

        excerpt = evidence["education"][0]["excerpt"]
        self.assertIn("Bachelor of Science in Information Technology", excerpt)
        self.assertIn("State University", excerpt)
        self.assertIn("2018 - 2022", excerpt)
        self.assertNotIn("Dean's Lister", excerpt)
        self.assertNotIn("Technical Support", excerpt)

    def test_generic_job_title_does_not_attach_to_neighboring_record(self):
        sample = """
Teaching Experience
Pampanga Colleges
2014 - 2018 Junior and Senior High School Teacher
San Miguel Academy, Incorporated
August 2022 - March 2026
"""
        experience = [SimpleNamespace(
            job_title="Teacher", company="San Miguel Academy",
            location="Pampanga", years=3.5,
        )]
        evidence = build_candidate_evidence(
            resume_text=sample, matched_skills=[], missing_skills=[],
            matched_critical_skills=[], missing_critical_skills=[],
            matched_preferred_skills=[], experience_records=experience,
            education_records=[],
        )

        excerpt = evidence["experience"][0]["excerpt"]
        self.assertIn("San Miguel Academy", excerpt)
        self.assertIn("August 2022 - March 2026", excerpt)
        self.assertNotIn("Pampanga Colleges", excerpt)
        self.assertNotIn("2014 - 2018", excerpt)

    def test_annotated_education_date_keeps_date_but_removes_award(self):
        sample = """
Educational Background
2022 - 2026 Tertiary Dean's Lister
Bachelor of Secondary Education major in English
Pampanga State University
2016 - 2022 Secondary With Highest Honors
Sto. Rosario National High School
"""
        education = [SimpleNamespace(
            degree="Bachelor of Secondary Education major in English",
            institution="Pampanga State University",
        )]
        evidence = build_candidate_evidence(
            resume_text=sample, matched_skills=[], missing_skills=[],
            matched_critical_skills=[], missing_critical_skills=[],
            matched_preferred_skills=[], experience_records=[],
            education_records=education,
        )

        excerpt = evidence["education"][0]["excerpt"]
        self.assertIn("2022 - 2026", excerpt)
        self.assertNotIn("Dean's Lister", excerpt)
        self.assertNotIn("Highest Honors", excerpt)


class TestExtractEducation(unittest.TestCase):

    def test_common_philippine_degree_abbreviation_is_recognized(self):
        sample = """
Education
BSEd major in Mathematics
Institution: Philippine Normal University
2018 - 2022
"""
        records = extract_education(sample)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["degree"], "BSEd major in Mathematics")
        self.assertEqual(records[0]["institution"], "Philippine Normal University")

    def test_inline_education_heading_extracts_degree_and_school(self):
        sample = """
Senior Agile Program Manager and Certified Scrum Master
SUMMARY:
Experienced in product management and group project delivery.
EDUCATION: Bachelor of Engineering (Electronics & Communication) 1998, .Karnataka University India
PROFESSIONAL TRAINING AND CERTIFICATIONS:
PMP, Scrum Professional, and Java certification
PROFESSIONAL EXPERIENCE:
Group Project Manager
"""
        records = extract_education(sample)
        self.assertEqual(records, [{
            "degree": "Bachelor of Engineering (Electronics & Communication)",
            "institution": "Karnataka University India",
        }])

    def test_scrum_master_is_not_treated_as_masters_degree(self):
        sample = """
Professional Summary
Certified Scrum Master with ten years of project management experience.
Professional Experience
Scrum Master
ABC Corporation
2018 - Present
"""
        self.assertEqual(extract_education(sample), [])

    def test_school_before_degree_and_multiple_entries_are_supported(self):
        sample = """
Academic Background
Philippine Normal University
Master of Arts in Education, 2022

University of Santo Tomas
BSEd major in English, 2018
Professional Experience
English Teacher
"""
        records = extract_education(sample)
        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["degree"], "Master of Arts in Education")
        self.assertEqual(records[0]["institution"], "Philippine Normal University")
        self.assertEqual(records[1]["degree"], "BSEd major in English")
        self.assertEqual(records[1]["institution"], "University of Santo Tomas")

    def test_degree_does_not_borrow_school_across_blank_entry_boundary(self):
        sample = """
Education
Bachelor of Science in Information Technology

Central State University
Master of Arts in Education
"""
        records = extract_education(sample)

        self.assertEqual(records[0]["institution"], "Unknown Institution")
        self.assertEqual(records[1]["institution"], "Central State University")

    def test_uppercase_institution_and_address_are_normalized(self):
        sample = """
EDUCATIONAL BACKGROUND
BACHELOR OF SECONDARY EDUCATION MAJOR IN ENGLISH
PAMPANGA STATE UNIVERSITY, BACOLOR, PAMPANGA
2022 - 2026
"""
        records = extract_education(sample)

        self.assertEqual(records[0]["institution"], "PAMPANGA STATE UNIVERSITY")

    def test_labelled_institution_acronym_is_supported(self):
        sample = """
Education
Bachelor of Science in Information Technology
Institution: PUP
2019 - 2023
"""
        records = extract_education(sample)

        self.assertEqual(records[0]["institution"], "PUP")

    def test_inline_degree_and_institution_do_not_merge_fields(self):
        sample = """
Education
Bachelor of Science in Computer Science | Polytechnic University of the Philippines
2018 - 2022
"""
        records = extract_education(sample)

        self.assertEqual(
            records[0]["institution"],
            "Polytechnic University of the Philippines",
        )

    def test_company_in_education_section_is_not_an_institution(self):
        sample = """
Education
Bachelor of Science in Computer Science
Training sponsored by Acme Corporation
Work Experience
Software Engineer
"""
        records = extract_education(sample)

        self.assertEqual(records[0]["institution"], "Unknown Institution")

    def test_academic_sentence_is_not_mistaken_for_education_heading(self):
        sample = """
Professional Summary
Committed to the academic growth of learners.
Work Experience
English Teacher
Education
Bachelor of Secondary Education
Pampanga Colleges
"""
        records = extract_education(sample)

        self.assertEqual(records[0]["institution"], "Pampanga Colleges")


class TestCertificationPrecision(unittest.TestCase):

    def test_combined_heading_and_inline_certification_list(self):
        sample = """
CERTIFICATIONS, SKILLS & AWARDS
Skills: Customer Service; Communication; Teamwork
Certifications: Certificate of Completion in Introduction to Food and Beverage Services; Standard First Aid and BLS CPR/AED Training; Microcertificate of Completion in Customer Centricity
Awards: Graduated Senior High School With High Honors
"""
        records = extract_certifications(sample)

        self.assertEqual([record["certification_name"] for record in records], [
            "Certificate of Completion in Introduction to Food and Beverage Services",
            "Standard First Aid and BLS CPR/AED Training",
            "Microcertificate of Completion in Customer Centricity",
        ])

    def test_french_languages_section_stops_certification_capture(self):
        sample = """
CERTIFICATIONS
CISCO 2025
Cyber Sécurité niveau 1
LANGUES
Français
Anglais
AHOUNOU Evrard
"""
        records = extract_certifications(sample)

        self.assertEqual(
            [record["certification_name"] for record in records],
            ["CISCO", "Cyber Sécurité niveau 1"],
        )

    def test_certification_section_stops_at_uncommon_resume_section(self):
        sample = """
Certifications
AWS Certified Cloud Practitioner - 2024
References
Maria Reyes, HR Manager
Available upon request
"""
        records = extract_certifications(sample)

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["certification_name"], "AWS Certified Cloud Practitioner")

    def test_certification_does_not_absorb_experience_narrative(self):
        sample = """
Professional Summary
Certified Scrum Master with 10 years of project management experience.
"""
        records = extract_certifications(sample)

        self.assertEqual(records[0]["certification_name"], "Certified Scrum Master")


class TestDocxExtraction(unittest.TestCase):

    def test_footer_contact_details_are_extracted(self):
        import docx

        document = docx.Document()
        document.add_paragraph("Professional Summary")
        document.sections[0].footer.paragraphs[0].text = "maria@example.com"

        with tempfile.TemporaryDirectory() as temp_dir:
            path = os.path.join(temp_dir, "footer_resume.docx")
            document.save(path)
            extracted = extract_text_from_docx(path)

        self.assertIn("maria@example.com", extracted)

    def test_text_box_content_is_extracted(self):
        import docx
        from docx.oxml import OxmlElement

        document = docx.Document()
        paragraph = document.add_paragraph("Skills")
        textbox = OxmlElement("w:txbxContent")
        textbox_paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "Python and Data Analysis"
        run.append(text)
        textbox_paragraph.append(run)
        textbox.append(textbox_paragraph)
        paragraph._p.append(textbox)

        with tempfile.TemporaryDirectory() as temp_dir:
            path = os.path.join(temp_dir, "textbox_resume.docx")
            document.save(path)
            extracted = extract_text_from_docx(path)

        self.assertIn("Python and Data Analysis", extracted)

    def test_pdf_quality_selector_prefers_structured_text(self):
        structured = """MARIA SANTOS
Professional Summary
Experienced software engineer
Work Experience
Software Engineer
Acme Systems
2020 - Present
Education
Bachelor of Science in Computer Science
Skills
Python SQL"""
        flattened = structured.replace("\n", " ")

        selected, method = _select_best_pdf_text([
            ("flattened", flattened),
            ("structured", structured),
        ])

        self.assertEqual(method, "structured")
        self.assertEqual(selected, structured)

    def test_pdf_quality_penalizes_concatenated_headings(self):
        readable = "Work Experience\nSoftware Engineer\nAcme Systems\n2020 - 2024"
        concatenated = "WORKEXPERIENCE\nSoftware Engineer\nAcme Systems\n2020 - 2024"

        self.assertGreater(
            _text_quality_score(readable),
            _text_quality_score(concatenated),
        )

    def test_cleaner_removes_adjacent_duplicate_lines(self):
        raw = "Skills\nClassroom Management\n  classroom   management  \nEducation"
        cleaned = _clean_extracted_text(raw)

        self.assertEqual(cleaned, "Skills\nClassroom Management\nEducation")

    def test_header_and_body_tables_preserve_reading_order(self):
        import docx

        document = docx.Document()
        document.sections[0].header.paragraphs[0].text = "MARIA SANTOS"
        document.add_paragraph("Work Experience")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Teacher"
        table.cell(0, 1).text = "San Isidro School"
        document.add_paragraph("June 2020 - Present")

        with tempfile.TemporaryDirectory() as temp_dir:
            path = os.path.join(temp_dir, "ordered_resume.docx")
            document.save(path)
            extracted = extract_text_from_docx(path)

        self.assertLess(extracted.index("MARIA SANTOS"), extracted.index("Work Experience"))
        self.assertLess(extracted.index("Work Experience"), extracted.index("Teacher"))
        self.assertLess(extracted.index("Teacher"), extracted.index("June 2020 - Present"))

    def test_vertically_merged_resume_table_does_not_repeat_into_skills(self):
        import docx

        document = docx.Document()
        table = document.add_table(rows=6, cols=2)
        table.cell(0, 1).text = "ARENAS, Rhea Hilary A."
        main_content = table.cell(1, 1).merge(table.cell(5, 1))
        main_content.text = """Experience
Staff Writer
CAST Chronicle
2019-2021
Education
Bachelor of Secondary Education major in English
Pangasinan State University
2018-2022
Skills
Computer Literate
Leadership
Patience
Verbal or Written Communication
Lesson Planning"""
        table.cell(2, 0).text = "Purok 4, Barangay Anunas, Angeles City, Pampanga"
        table.cell(3, 0).text = "0956-766-9914"
        table.cell(4, 0).text = "arenasrheahilary@gmail.com"
        table.cell(5, 0).text = "References\nBernie O. Pioquinto"

        with tempfile.TemporaryDirectory() as temp_dir:
            path = os.path.join(temp_dir, "vertical_merge_resume.docx")
            document.save(path)
            extracted = extract_text_from_docx(path)

        self.assertEqual(extracted.count("Experience"), 1)
        self.assertEqual(extracted.count("Skills"), 1)
        self.assertEqual(extract_resume_skills(extracted), [
            "Computer Literate",
            "Leadership",
            "Patience",
            "Verbal or Written Communication",
            "Lesson Planning",
        ])

    def test_cleaner_preserves_wrapped_date_ranges(self):
        raw = "Work Experience\nSoftware Engineer\nAcme Digital Solutions\nJanuary 2020 -\nPresent"
        cleaned = _clean_extracted_text(raw)

        self.assertIn("January 2020 - Present", cleaned)

    def test_cleaner_still_joins_soft_hyphenated_words(self):
        raw = "Professional Summary\nDevelop-\nment of internal applications"
        cleaned = _clean_extracted_text(raw)

        self.assertIn("Development of internal applications", cleaned)

    def test_cleaner_preserves_role_company_separator_before_capitalized_employer(self):
        cleaned = _clean_extracted_text(
            "Work Experience\nActivities Coordinator-\nNational University\n2023-2024"
        )

        self.assertIn("Activities Coordinator-\nNational University", cleaned)

    def test_cleaner_removes_invisible_word_break_artifacts(self):
        cleaned = _clean_extracted_text("Data\u200bbase Admin\u00adistrator")

        self.assertEqual(cleaned, "Database Administrator")

    def test_cleaner_normalizes_private_font_bullets(self):
        cleaned = _clean_extracted_text("Skills\n\uf06c  Classroom Management")

        self.assertEqual(cleaned, "Skills\n- Classroom Management")

    def test_pdf_and_docx_style_date_wrapping_extract_same_experience(self):
        pdf_style = _clean_extracted_text("""
Work Experience
Software Engineer
Acme Digital Solutions
January 2020 -
Present
""")
        docx_style = _clean_extracted_text("""
Work Experience
Software Engineer
Acme Digital Solutions
January 2020 - Present
""")

        pdf_records = extract_experience_records(pdf_style)
        docx_records = extract_experience_records(docx_style)

        self.assertEqual(pdf_records, docx_records)
        self.assertEqual(pdf_records[0]["job_title"], "Software Engineer")
        self.assertEqual(pdf_records[0]["company"], "Acme Digital Solutions")

    def test_two_column_pdf_keeps_education_institution_in_education_column(self):
        class PositionedPage:
            width = 600

            def __init__(self):
                self.words = []

            def add_line(self, x, top, value):
                cursor = x
                for token in value.split():
                    width = max(len(token) * 5, 8)
                    self.words.append({
                        "text": token, "x0": cursor, "x1": cursor + width,
                        "top": top,
                    })
                    cursor += width + 4

            def extract_words(self, **_kwargs):
                return self.words

        page = PositionedPage()
        # Right column starts higher, as in templates with a photo/sidebar.
        for index, value in enumerate([
            "MARIA SANTOS", "PROFESSIONAL SUMMARY",
            "Experienced classroom educator", "TEACHING EXPERIENCE",
            "English Teacher", "Bright Future Academy", "2022 - Present",
            "Prepared lessons and assessments for students",
        ]):
            page.add_line(250, 20 + index * 24, value)
        for index, value in enumerate([
            "EDUCATION", "Bachelor of Secondary Education", "Major in English",
            "Pampanga State University", "2018 - 2022", "SKILLS",
            "Curriculum Planning", "Classroom Management", "Communication",
            "Leadership", "Assessment", "Lesson Planning",
        ]):
            page.add_line(25, 170 + index * 24, value)

        extracted = _extract_page_columns(page)

        self.assertLess(extracted.index("PROFESSIONAL SUMMARY"), extracted.index("TEACHING EXPERIENCE"))
        self.assertLess(extracted.index("EDUCATION"), extracted.index("Pampanga State University"))
        education = extract_education(extracted)
        self.assertEqual(education[0]["institution"], "Pampanga State University")

    def test_crossing_name_two_column_pdf_reconstructs_degree_and_work_history(self):
        class PositionedPage:
            width = 600

            def __init__(self):
                self.words = [
                    {"text": "MELVIN", "x0": 215, "x1": 400, "top": 20},
                    {"text": "AGUILAR", "x0": 215, "x1": 420, "top": 65},
                ]

            def add_line(self, x, top, value):
                cursor = x
                for token in value.split():
                    width = max(len(token) * 4.2, 7)
                    self.words.append({
                        "text": token,
                        "x0": cursor,
                        "x1": cursor + width,
                        "top": top,
                    })
                    cursor += width + 3

            def extract_words(self, **_kwargs):
                return self.words

        page = PositionedPage()
        for index, value in enumerate([
            "EXPERIENCES",
            "Student Development and Activities Coordinator-",
            "National University Clark",
            "2023-2024",
            "Professional Event Host/ Event Manager/ Event Planner",
            "and Coordinator- Freelance",
            "2016- Present",
            "Full Time Faculty Member- National University Clark",
            "2023-2024",
            "Full Time Faculty Member- University of the Assumption",
            "2019-2023",
            "Part Time College Instructor- City College of San",
            "Fernando Pampanga",
            "2021-2023",
            "EDUCATION",
            "City College of San Fernando Pampanga",
            "Bachelor of Secondary Education Major in Biological",
            "Science",
            "2015-2019",
        ]):
            page.add_line(255, 125 + index * 18, value)
        for index, value in enumerate([
            "About Me", "Experienced educator and event host",
            "Teaching science subjects in senior high school",
            "Professional educator with classroom experience",
            "+63 968 597 6305", "melvinaguilar@example.com",
            "San Fernando Pampanga Philippines",
            "CERTIFICATES AND LICENSE", "Board passer",
        ]):
            page.add_line(25, 170 + index * 24, value)

        extracted = _clean_extracted_text(_extract_page_columns(page))

        self.assertEqual(extract_contact_info(extracted)["name"], "Melvin Aguilar")
        self.assertEqual(extract_education(extracted), [{
            "degree": "Bachelor of Secondary Education Major in Biological Science",
            "institution": "City College of San Fernando Pampanga",
        }])
        self.assertEqual([
            (record["job_title"], record["company"])
            for record in extract_experience_records(extracted)
        ], [
            (
                "Professional Event Host/ Event Manager/ Event Planner and Coordinator",
                "Freelance",
            ),
            ("Full Time Faculty Member", "National University Clark"),
            ("Full Time Faculty Member", "University of the Assumption"),
            ("Part Time College Instructor", "City College of San Fernando Pampanga"),
            ("Student Development and Activities Coordinator", "National University Clark"),
        ])


class TestAnalyzeSkills(unittest.TestCase):

    # --- word-boundary correctness ---

    def test_short_skill_not_matched_inside_longer_word(self):
        """'R' must not match inside 'recruiter', 'Go' not inside 'algorithm'."""
        matched, missing = analyze_skills(
            "Experienced recruiter with algorithm design exposure.",
            ["R", "Go", "algorithm"],
        )
        self.assertEqual(matched, ["algorithm"])
        self.assertEqual(missing, ["R", "Go"])

    def test_exact_short_skill_matched_standalone(self):
        """'R' and 'Go' should match when they appear as standalone tokens."""
        matched, missing = analyze_skills(
            "Proficient in R and Go for data pipelines.",
            ["R", "Go"],
        )
        self.assertIn("R", matched)
        self.assertIn("Go", matched)
        self.assertEqual(missing, [])

    # --- alias resolution ---

    def test_aliases_resolve_javascript_variants(self):
        """'JavaScript' should match when resume contains 'JS'."""
        matched, missing = analyze_skills("Built SPAs with JS.", ["JavaScript"])
        self.assertEqual(matched, ["JavaScript"])
        self.assertEqual(missing, [])

    def test_aliases_resolve_full_stack_common_variants(self):
        """CSS3, HTML5, React.js, PostgreSQL should resolve to their canonical names."""
        resume = "Built dashboards with JS, React.js, HTML5, CSS3, and PostgreSQL."
        skills = ["JavaScript", "React", "HTML", "CSS", "Postgres"]
        matched, missing = analyze_skills(resume, skills)
        self.assertEqual(sorted(matched), sorted(skills))
        self.assertEqual(missing, [])

    def test_nlp_alias_natural_language_processing(self):
        """'Natural Language Processing' should match 'NLP' in resume text."""
        matched, missing = analyze_skills(
            "Worked on NLP tasks for sentiment analysis.",
            ["Natural Language Processing"],
        )
        self.assertEqual(matched, ["Natural Language Processing"])

    def test_python_versioned_alias(self):
        """'Python' should match 'Python 3' in the resume."""
        matched, missing = analyze_skills(
            "Developed REST APIs with Python 3 and Flask.",
            ["Python"],
        )
        self.assertEqual(matched, ["Python"])

    # --- edge cases ---

    def test_empty_resume_text_yields_all_missing(self):
        matched, missing = analyze_skills("", ["Python", "SQL"])
        self.assertEqual(matched, [])
        self.assertEqual(missing, ["Python", "SQL"])

    def test_empty_required_skills_returns_nothing(self):
        matched, missing = analyze_skills("Experienced developer.", [])
        self.assertEqual(matched, [])
        self.assertEqual(missing, [])

    def test_blank_skill_entries_are_skipped(self):
        """Blank strings in the required list should be silently ignored."""
        matched, missing = analyze_skills("Knows Python.", ["Python", "", "  "])
        self.assertEqual(matched, ["Python"])
        self.assertEqual(missing, [])

    def test_case_insensitive_matching(self):
        matched, missing = analyze_skills("PYTHON and SQL expert.", ["python", "sql"])
        self.assertEqual(sorted(matched), ["python", "sql"])
        self.assertEqual(missing, [])

    def test_special_character_skill_csharp(self):
        """'C#' should match literally (no word-boundary regex applied to it)."""
        matched, _ = analyze_skills("Proficient in C# and .NET.", ["C#"])
        self.assertIn("C#", matched)


class TestExtractResumeSkills(unittest.TestCase):

    def test_extracts_wrapped_inline_skills_from_combined_section(self):
        skills = extract_resume_skills("""
CERTIFICATIONS, SKILLS & AWARDS
▪ Skills: Customer Service; Communication; Teamwork; Problem-Solving; Attention to
Detail; Event Coordination; Sanitation Practices
▪ Certifications: Standard First Aid
▪ Awards: High Honors
""")

        self.assertEqual(skills, [
            "Customer Service", "Communication", "Teamwork", "Problem-Solving",
            "Attention to Detail", "Event Coordination", "Sanitation Practices",
        ])

    def test_richer_explicit_communication_phrase_blocks_generic_duplicate(self):
        skills = extract_resume_skills("""
Skills
Communicates clearly in written and verbal contexts
Education
Bachelor of Education
Work Experience
Communicated with parents and students daily.
""")

        self.assertIn("Communicates clearly in written and verbal contexts", skills)
        self.assertNotIn("communication skills", [skill.lower() for skill in skills])

    def test_generic_skill_suffix_does_not_create_duplicate(self):
        skills = extract_resume_skills("""
Skills
Problem Solving
Problem-Solving Skills
Education
Bachelor of Education
""")

        self.assertEqual(skills, ["Problem Solving"])

    def test_extracts_skills_not_required_by_job(self):
        resume = """
Technical Skills
Python, Flask, Docker, Git
Education
Bachelor of Science in Computer Science
"""
        skills = extract_resume_skills(resume, additional_skills=["Python", "SQL"])

        self.assertEqual(skills, ["Python", "Flask", "Docker", "Git"])
        self.assertNotIn("SQL", skills)

    def test_preserves_explicit_resume_wording_instead_of_alias_names(self):
        resume = """
Skills: JS, JavaScript, PostgreSQL, Postgres, MS Excel
"""
        skills = extract_resume_skills(resume)

        self.assertEqual(
            skills,
            ["JS", "JavaScript", "PostgreSQL", "Postgres", "MS Excel"],
        )

    def test_preserves_long_competency_phrase_from_skills_section(self):
        phrase = "Ability to organize, prioritize, and manage multiple classroom tasks"
        communication = "Written and oral communication skills"
        skills = extract_resume_skills(
            f"Core Competencies\n{phrase}\n{communication}\nEducation"
        )

        self.assertIn(phrase, skills)
        self.assertIn(communication, skills)
        self.assertNotIn("communication skills", skills)

    def test_splits_pdf_concatenated_competency_phrases(self):
        skills = extract_resume_skills("""
Skills
Designs engaging lesson plans Communicates clearly with parents
Leverages online platforms Adapts quickly to changing situations
Education
Bachelor of Education
""")

        self.assertEqual(skills, [
            "Designs engaging lesson plans",
            "Communicates clearly with parents",
            "Leverages online platforms",
            "Adapts quickly to changing situations",
        ])

    def test_category_labels_are_removed_and_parenthetical_commas_are_kept(self):
        skills = extract_resume_skills("""
Technical Skills
Cloud Technologies  Amazon Web Services (EC2, SQS, RDS, IAM, S3)
Frontend  JavaScript, HTML, CSS
IDE s: Eclipse, IntelliJ IDEA
Education
Bachelor of Science
""")

        self.assertEqual(skills, [
            "Amazon Web Services (EC2, SQS, RDS, IAM, S3)",
            "JavaScript",
            "HTML",
            "CSS",
            "Eclipse",
            "IntelliJ IDEA",
        ])

    def test_inline_education_heading_stops_skill_section(self):
        skills = extract_resume_skills("""
Technical Skills: Python, SQL
EDUCATION: Bachelor of Science, State University
PROFESSIONAL TRAINING AND CERTIFICATIONS: AWS Practitioner
""")

        self.assertEqual(skills, ["Python", "SQL"])

    def test_structural_separator_does_not_turn_following_list_into_one_skill(self):
        skills = extract_resume_skills("""
Technical Skills
Software Development Life Cycle (SDLC) - Agile, SCRUM, Waterfall
Education
Bachelor of Science
""")

        self.assertEqual(skills, [
            "Software Development Life Cycle (SDLC)",
            "Agile",
            "SCRUM",
            "Waterfall",
        ])

    def test_wide_pdf_columns_keep_first_skill_and_technology_year(self):
        skills = extract_resume_skills("""
Technical Skills
Python    Docker
Operating Systems  Red Hat Linux 9, Unix/Linux, Windows 2000/NT/XP
Education
Bachelor of Science
""")

        self.assertEqual(skills, [
            "Python",
            "Docker",
            "Red Hat Linux 9",
            "Unix/Linux",
            "Windows 2000/NT/XP",
        ])

    def test_very_long_explicit_competency_is_not_dropped(self):
        phrase = (
            "Ability to design inclusive learning activities and communicate "
            "complex instructions clearly while coordinating with learners, "
            "parents, teachers, administrators, and community stakeholders"
        )
        skills = extract_resume_skills(f"Skills\n{phrase}\nEducation")

        self.assertEqual(skills, [phrase])

    def test_skill_phrases_are_not_cut_at_action_words_or_product_hyphens(self):
        skills = extract_resume_skills("""
Skills & Qualities
  Develop Lesson Plans
Testing: Cucumber - JVM
Skilful in Conflict Resolution, Risk Identification and Mitigation
Education
Bachelor of Education
""")

        self.assertEqual(skills, [
            "Develop Lesson Plans",
            "Testing",
            "Cucumber - JVM",
            "Skilful in Conflict Resolution, Risk Identification and Mitigation",
        ])

    def test_meaningful_labels_are_kept_but_category_labels_are_removed(self):
        skills = extract_resume_skills("""
Technical Skills
Spring Framework: Spring Boot, Spring Security
NoSQL: Cassandra, MongoDB
Design Methodologies: UML, OOAD, Design Patterns
IDEs / Tools  Eclipse, IntelliJ IDEA
Project Management tools  JIRA, Trello, and SharePoint
Education
Bachelor of Science
""")

        self.assertEqual(skills, [
            "Spring Framework",
            "Spring Boot",
            "Spring Security",
            "NoSQL",
            "Cassandra",
            "MongoDB",
            "UML",
            "OOAD",
            "Design Patterns",
            "Eclipse",
            "IntelliJ IDEA",
            "JIRA",
            "Trello",
            "SharePoint",
        ])

    def test_pdf_wrapped_bullet_skills_are_reconstructed_before_splitting(self):
        skills = extract_resume_skills("""
SKILLS & QUALITIES
  Curriculum and Subject
Matter knowledge
  Verbal and Written
Communication Skills
  Leadership
Curriculum Implementation
  Modifying Instructional
Materials
  Integrating Technology in
Teaching and Learning
Processes
  Utilizing Various Teaching
Strategies
  Providing Interactive
Activities/Game
Curriculum Evaluation
  Excellent in Formative and
Summative Assessment
Education
Bachelor of Education
""")

        self.assertEqual(skills, [
            "Curriculum and Subject Matter knowledge",
            "Verbal and Written Communication Skills",
            "Leadership",
            "Curriculum Implementation",
            "Modifying Instructional Materials",
            "Integrating Technology in Teaching and Learning Processes",
            "Utilizing Various Teaching Strategies",
            "Providing Interactive Activities/Game",
            "Curriculum Evaluation",
            "Excellent in Formative and Summative Assessment",
        ])

    def test_combined_skills_qualities_and_interests_headings_are_supported(self):
        qualities = extract_resume_skills("""
Skills & Qualities
Patient and dependable
Clear written communication
Education
Bachelor of Education
""")
        interests = extract_resume_skills("""
SKILLS / INTERESTS: Event Hosting, Theater Acting
Work Experience
Teacher
""")

        self.assertEqual(qualities, [
            "Patient and dependable",
            "Clear written communication",
        ])
        self.assertEqual(interests, ["Event Hosting", "Theater Acting"])

    def test_explicit_skill_section_is_authoritative_for_display_inventory(self):
        skills = extract_resume_skills("""
Skills
Python, Clear written communication
Work Experience
Built and deployed services using Docker and Kubernetes.
Education
Bachelor of Science in Information Technology
""")

        self.assertEqual(skills, ["Python", "Clear written communication"])

    def test_contextual_alias_is_not_duplicated_when_explicitly_listed(self):
        skills = extract_resume_skills("""
Skills: JS, Postgres
Projects
Created JavaScript dashboards backed by PostgreSQL.
""", additional_skills=["JS", "PostgreSQL"])

        self.assertEqual(skills, ["JS", "Postgres"])

    def test_job_configured_alias_does_not_create_case_variant_duplicate(self):
        skills = extract_resume_skills("""
Skills: Communication Skills
Work Experience
Demonstrated strong communication skills when assisting clients.
""", additional_skills=["Communication Skills"])

        self.assertEqual(skills, ["Communication Skills"])

    def test_negated_and_aspirational_context_skills_are_not_extracted(self):
        skills = extract_resume_skills("""
Professional Summary
No experience with Docker. Interested in learning Kubernetes.
Work Experience
Developed Python services with Flask.
""")

        self.assertIn("Python", skills)
        self.assertIn("Flask", skills)
        self.assertNotIn("Docker", skills)
        self.assertNotIn("Kubernetes", skills)

    def test_education_only_technology_is_not_contextual_skill_evidence(self):
        skills = extract_resume_skills("""
Education
Bachelor of Science in Information Technology
Python Institute Certificate Program
References
Available upon request
""")

        self.assertNotIn("Python", skills)

    def test_accepts_unknown_explicit_skill_but_not_resume_prose(self):
        resume = """
Core Competencies
Records Classification
Work Experience
Managed records for Acme Corporation and prepared weekly reports.
"""
        skills = extract_resume_skills(resume)

        self.assertIn("Records Classification", skills)
        self.assertNotIn("Managed Records For Acme Corporation And Prepared Weekly Reports", skills)

    def test_does_not_extract_negated_catalog_skill(self):
        skills = extract_resume_skills(
            "Professional Summary\nNo experience with Docker.\nSkills: Without Kubernetes"
        )

        self.assertNotIn("Docker", skills)
        self.assertNotIn("Kubernetes", skills)


# ===========================================================================
# analyze_preferred_skills
# ===========================================================================

class TestAnalyzePreferredSkills(unittest.TestCase):

    def test_bonus_capped_at_10_points_with_full_match(self):
        """When ALL preferred skills match, bonus should be exactly 10.0."""
        matched, total, bonus = analyze_preferred_skills(
            "Expert in Docker, Kubernetes, and Redis.",
            ["Docker", "Kubernetes", "Redis"],
        )
        self.assertEqual(len(matched), 3)
        self.assertEqual(total, 3)
        self.assertEqual(bonus, 10.0)

    def test_partial_preferred_match_scales_proportionally(self):
        """2 out of 3 preferred → bonus = round(2/3 * 10, 2) = 6.67."""
        matched, total, bonus = analyze_preferred_skills(
            "Experience with NLP and Python 3.",
            ["Natural Language Processing", "Python", "Docker"],
        )
        self.assertEqual(matched, ["Natural Language Processing", "Python"])
        self.assertEqual(total, 3)
        self.assertEqual(bonus, 6.67)

    def test_empty_preferred_list_returns_zero(self):
        matched, total, bonus = analyze_preferred_skills("Some resume text.", [])
        self.assertEqual(matched, [])
        self.assertEqual(total, 0)
        self.assertEqual(bonus, 0.0)

    def test_no_preferred_skills_match(self):
        matched, total, bonus = analyze_preferred_skills(
            "Knows only MS Word.",
            ["Kubernetes", "Terraform"],
        )
        self.assertEqual(matched, [])
        self.assertEqual(total, 2)
        self.assertEqual(bonus, 0.0)

    def test_blank_preferred_entries_are_ignored(self):
        """Blank entries must not inflate the total denominator."""
        matched, total, bonus = analyze_preferred_skills(
            "Expert in Docker.",
            ["Docker", "", "  "],
        )
        self.assertEqual(total, 1)
        self.assertEqual(bonus, 10.0)


# ===========================================================================
# generate_recommendation
# ===========================================================================

class TestGenerateRecommendation(unittest.TestCase):

    def test_qualified_at_exactly_75(self):
        self.assertEqual(generate_recommendation(75.0, 50), "Qualified")

    def test_qualified_above_75(self):
        self.assertEqual(generate_recommendation(99.0, 50), "Qualified")

    def test_for_review_at_min_fit_score(self):
        self.assertEqual(generate_recommendation(50.0, 50), "For Review")

    def test_for_review_between_min_and_75(self):
        self.assertEqual(generate_recommendation(65.0, 50), "For Review")

    def test_not_qualified_just_below_min(self):
        self.assertEqual(generate_recommendation(49.99, 50), "Not Qualified")

    def test_not_qualified_at_zero(self):
        self.assertEqual(generate_recommendation(0.0, 50), "Not Qualified")

    def test_custom_min_fit_score(self):
        """min_fit_score of 60 should gate the thresholds correctly."""
        self.assertEqual(generate_recommendation(60.0, 60), "For Review")
        self.assertEqual(generate_recommendation(59.9, 60), "Not Qualified")

    def test_boundary_at_75_exactly(self):
        """74.99 should be 'For Review', 75.0 should be 'Qualified'."""
        self.assertEqual(generate_recommendation(74.99, 50), "For Review")
        self.assertEqual(generate_recommendation(75.0, 50), "Qualified")


# ===========================================================================
# get_degree_rank
# ===========================================================================

class TestGetDegreeRank(unittest.TestCase):

    def test_phd_returns_5(self):
        self.assertEqual(get_degree_rank("Ph.D."), 5)
        self.assertEqual(get_degree_rank("Doctor of Philosophy"), 5)

    def test_masters_returns_4(self):
        self.assertEqual(get_degree_rank("Master's Degree"), 4)
        self.assertEqual(get_degree_rank("MBA"), 4)
        self.assertEqual(get_degree_rank("M.S."), 4)

    def test_bachelors_returns_3(self):
        self.assertEqual(get_degree_rank("Bachelor of Science"), 3)
        self.assertEqual(get_degree_rank("B.S."), 3)
        self.assertEqual(get_degree_rank("BSEd major in Mathematics"), 3)

    def test_associate_returns_2(self):
        self.assertEqual(get_degree_rank("Associate's Degree"), 2)

    def test_high_school_returns_1(self):
        self.assertEqual(get_degree_rank("High School Diploma"), 1)
        self.assertEqual(get_degree_rank("GED"), 1)

    def test_unknown_degree_returns_0(self):
        self.assertEqual(get_degree_rank("Certification"), 0)
        self.assertEqual(get_degree_rank(""), 0)
        self.assertIsNone(get_degree_rank(None)) if False else self.assertEqual(get_degree_rank(None), 0)


# ===========================================================================
# calculate_fit_score
# ===========================================================================

class TestCalculateFitScore(unittest.TestCase):

    def test_default_weights_50_30_20(self):
        """skill=80, exp=50, edu=100 → 80*.5 + 50*.3 + 100*.2 = 40+15+20 = 75"""
        self.assertEqual(calculate_fit_score(80, 50, 100), 75.0)

    def test_result_capped_at_100(self):
        self.assertEqual(calculate_fit_score(100, 100, 100), 100.0)

    def test_all_zero_inputs(self):
        self.assertEqual(calculate_fit_score(0, 0, 0), 0.0)

    def test_custom_weights(self):
        """Equal weights (1/3 each) with equal inputs should stay the same."""
        score = calculate_fit_score(60, 60, 60, weights=(1/3, 1/3, 1/3))
        self.assertAlmostEqual(score, 60.0, places=1)

    def test_returns_rounded_two_decimal_places(self):
        score = calculate_fit_score(33, 33, 33)
        # 33*.5 + 33*.3 + 33*.2 = 16.5 + 9.9 + 6.6 = 33.0 (exact)
        self.assertEqual(score, 33.0)


# ===========================================================================
# calculate_text_similarity
# ===========================================================================

class TestCalculateTextSimilarity(unittest.TestCase):

    def test_identical_texts_return_100(self):
        score = calculate_text_similarity("python machine learning", "python machine learning")
        self.assertEqual(score, 100.0)

    def test_completely_different_texts_return_0(self):
        score = calculate_text_similarity("python data science", "cooking recipes baking")
        self.assertEqual(score, 0.0)

    def test_empty_text1_returns_0(self):
        self.assertEqual(calculate_text_similarity("", "some text"), 0.0)

    def test_empty_text2_returns_0(self):
        self.assertEqual(calculate_text_similarity("some text", ""), 0.0)

    def test_both_empty_returns_0(self):
        self.assertEqual(calculate_text_similarity("", ""), 0.0)

    def test_similar_texts_score_higher_than_dissimilar(self):
        similar = calculate_text_similarity(
            "software engineer python sql", "developer python sql databases"
        )
        dissimilar = calculate_text_similarity(
            "software engineer python sql", "nurse hospital medicine surgery"
        )
        self.assertGreater(similar, dissimilar)


# ===========================================================================
# calculate_skills_match
# ===========================================================================

class TestCalculateSkillsMatch(unittest.TestCase):

    def test_full_match_returns_100(self):
        score = calculate_skills_match(["Python", "SQL"], ["Python", "SQL"])
        self.assertEqual(score, 100.0)

    def test_no_match_returns_0(self):
        score = calculate_skills_match(["Java"], ["Python", "SQL"])
        self.assertEqual(score, 0.0)

    def test_partial_match(self):
        score = calculate_skills_match(["Python", "Java"], ["Python", "SQL"])
        self.assertEqual(score, 50.0)

    def test_empty_required_skills_returns_100(self):
        """If a job requires no skills, everyone qualifies."""
        score = calculate_skills_match(["Python"], [])
        self.assertEqual(score, 100.0)

    def test_case_insensitive_matching(self):
        score = calculate_skills_match(["python", "sql"], ["Python", "SQL"])
        self.assertEqual(score, 100.0)


# ===========================================================================
# generate_analysis_narrative
# ===========================================================================

class TestGenerateAnalysisNarrative(unittest.TestCase):

    def test_strong_fit_opening_contains_strong_match(self):
        narrative = _make_narrative(fit_score=80.0)
        self.assertIn("strong match", narrative)

    def test_moderate_fit_opening(self):
        narrative = _make_narrative(fit_score=60.0)
        self.assertIn("moderate fit", narrative.lower())

    def test_weak_fit_opening(self):
        narrative = _make_narrative(
            fit_score=30.0, matched_skills=[], missing_skills=["Python", "SQL"]
        )
        self.assertIn("does not meet", narrative)

    def test_disqualified_by_critical_skills_overrides_tier(self):
        narrative = _make_narrative(
            fit_score=80.0, disqualified_by_critical_skills=True,
            missing_skills=["Python"]
        )
        self.assertIn("disqualified", narrative.lower())
        self.assertIn("Do not advance", narrative)

    def test_perfect_skill_match_says_all_skills(self):
        narrative = _make_narrative(
            matched_skills=["Python", "SQL"], missing_skills=[]
        )
        self.assertIn("all", narrative.lower())

    def test_zero_skills_match(self):
        narrative = _make_narrative(
            matched_skills=[], missing_skills=["Python", "SQL"],
            skill_score=0.0
        )
        self.assertIn("did not surface any", narrative)

    def test_preferred_skills_mentioned_when_matched(self):
        narrative = _make_narrative(
            matched_skills=["Python"],
            missing_skills=["SQL"],
            skill_score=60.0,
            matched_preferred=["Docker"],
            preferred_bonus=5.0,
        )
        self.assertIn("Docker", narrative)
        self.assertIn("+5pt bonus", narrative)

    def test_experience_meets_requirement(self):
        narrative = _make_narrative(total_exp_years=5.0, experience_req=3)
        self.assertIn("meeting or exceeding", narrative)

    def test_experience_falls_short(self):
        narrative = _make_narrative(total_exp_years=1.0, experience_req=3, exp_score=33.0)
        self.assertIn("falls short", narrative)

    def test_no_experience_req_set(self):
        narrative = _make_narrative(experience_req=0, total_exp_years=4.0)
        self.assertIn("No minimum experience was specified", narrative)

    def test_education_meets_requirement(self):
        narrative = _make_narrative(
            extracted_edu=[{"degree": "Bachelor's", "institution": "MIT"}],
            education_req="Bachelor's",
            edu_score=100.0,
        )
        self.assertIn("meets or exceeds", narrative)

    def test_education_does_not_meet_requirement(self):
        narrative = _make_narrative(
            extracted_edu=[{"degree": "High School Diploma", "institution": "Some School"}],
            education_req="Bachelor's",
            edu_score=33.0,
        )
        self.assertIn("does not meet", narrative)

    def test_no_education_requirement(self):
        narrative = _make_narrative(education_req=None)
        self.assertIn("No specific education requirement", narrative)

    def test_narrative_contains_job_title(self):
        narrative = _make_narrative(job_title="Data Scientist", fit_score=80.0)
        self.assertIn("Data Scientist", narrative)


class TestGenerateDecisionExplanation(unittest.TestCase):

    def test_explanation_makes_scoring_model_transparent(self):
        explanation = generate_decision_explanation(
            job_title="Software Engineer",
            recommendation_label="Qualified",
            fit_score=85.0,
            skill_score=90.0,
            exp_score=80.0,
            edu_score=100.0,
            text_similarity_score=42.0,
            matched_skills=["Python", "SQL"],
            missing_skills=["Docker"],
            matched_critical_skills=["Python"],
            missing_critical_skills=[],
            matched_preferred=["AWS"],
            preferred_bonus=5.0,
            total_exp_years=4.0,
            experience_req=3,
            education_req="Bachelor's",
        )
        self.assertIn(f"skills {FIT_WEIGHT_PERCENTS[0]}%", explanation)
        self.assertIn(f"experience {FIT_WEIGHT_PERCENTS[1]}%", explanation)
        self.assertIn(f"education {FIT_WEIGHT_PERCENTS[2]}%", explanation)
        self.assertIn("preferred skill", explanation.lower())
        self.assertIn(f"text similarity {FIT_WEIGHT_PERCENTS[3]}%", explanation)
        self.assertIn("matched", explanation.lower())
        self.assertIn("missing", explanation.lower())
        self.assertTrue(
            any(marker in explanation.lower() for marker in ["strength", "advantage", "works in"])
        )
        self.assertIn("work experience", explanation.lower())
        self.assertIn("advantage", explanation.lower())
        self.assertTrue(
            any(marker in explanation.lower() for marker in ["weakness", "concern", "gap"])
        )
        self.assertIn("Reviewer recommendation", explanation)

    def test_explanation_wording_varies_by_candidate_context(self):
        first = generate_decision_explanation(
            job_title="Software Engineer",
            recommendation_label="Qualified",
            fit_score=90.0,
            skill_score=100.0,
            exp_score=100.0,
            edu_score=100.0,
            text_similarity_score=40.0,
            matched_skills=["Python", "SQL"],
            missing_skills=[],
            matched_critical_skills=[],
            missing_critical_skills=[],
            matched_preferred=[],
            preferred_bonus=0.0,
            total_exp_years=5.0,
            experience_req=3,
            education_req="Bachelor's",
        )
        second = generate_decision_explanation(
            job_title="Guidance Counselor",
            recommendation_label="For Review",
            fit_score=68.0,
            skill_score=70.0,
            exp_score=60.0,
            edu_score=100.0,
            text_similarity_score=30.0,
            matched_skills=["Counseling"],
            missing_skills=["Case Management"],
            matched_critical_skills=[],
            missing_critical_skills=[],
            matched_preferred=[],
            preferred_bonus=0.0,
            total_exp_years=2.0,
            experience_req=3,
            education_req="Bachelor's",
        )
        self.assertNotEqual(first.split("\n\n")[0], second.split("\n\n")[0])


# ===========================================================================
# evaluate_candidate (integration / orchestration)
# ===========================================================================

class TestEvaluateCandidate(unittest.TestCase):

    RESUME = (
        "John Doe  |  john@example.com  |  +1-555-123-4567\n"
        "Summary: Senior software engineer with 5 years of experience.\n\n"
        "Skills: Python, SQL, JavaScript, React, Docker\n\n"
        "Education\n"
        "Bachelor of Science in Computer Science\n"
        "State University  2015 - 2019\n\n"
        "Experience\n"
        "Software Engineer at Acme Corp  2019 - 2024\n"
    )

    JOB_DESC = (
        "We are looking for a Software Engineer proficient in Python, SQL, "
        "and JavaScript to join our growing team."
    )

    def _run(self, **kwargs):
        defaults = dict(
            resume_text=self.RESUME,
            job_desc_text=self.JOB_DESC,
            required_skills=["Python", "SQL", "JavaScript"],
            min_fit_score=50.0,
            experience_req=3,
            education_req="Bachelor's",
            job_title="Software Engineer",
        )
        defaults.update(kwargs)
        return evaluate_candidate(**defaults)

    # --- return shape ---

    def test_result_contains_all_expected_keys(self):
        result = self._run()
        for key in [
            "skill_score", "experience_score", "education_score",
            "text_similarity_score", "fit_score", "recommendation_label",
            "confidence_level", "confidence_reason",
            "matched_skills", "missing_skills", "matched_preferred",
            "extracted_skills",
            "matched_critical_skills", "missing_critical_skills",
            "summary", "decision_explanation", "contact_info", "extracted_edu", "extracted_exp",
            "total_exp_years",
        ]:
            self.assertIn(key, result, msg=f"Missing key: {key}")

    # --- skill scoring ---

    def test_all_required_skills_matched(self):
        result = self._run()
        self.assertEqual(sorted(result["matched_skills"]),
                         sorted(["Python", "SQL", "JavaScript"]))

    def test_evaluation_returns_non_job_resume_skills_without_scoring_them(self):
        result = self._run()

        self.assertIn("Docker", result["extracted_skills"])
        self.assertNotIn("Docker", result["matched_skills"])
        self.assertEqual(result["missing_skills"], [])

    def test_skill_score_100_when_all_matched(self):
        result = self._run()
        self.assertEqual(result["skill_score"], 100.0)

    def test_missing_skills_reduce_skill_score(self):
        result = self._run(required_skills=["Python", "SQL", "Rust", "Haskell"])
        # Python + SQL matched, Rust + Haskell missing → 50%
        self.assertEqual(result["skill_score"], 50.0)
        self.assertIn("Rust", result["missing_skills"])

    # --- experience scoring ---

    def test_experience_score_100_when_exp_meets_requirement(self):
        result = self._run(experience_req=3)
        self.assertEqual(result["experience_score"], 100.0)

    def test_experience_score_partial_when_exp_below_requirement(self):
        """Candidate has 5 years; require 10 → raw_ratio 0.5 → exp_score 50."""
        result = self._run(experience_req=10)
        self.assertEqual(result["experience_score"], 50.0)

    def test_no_experience_requirement_gives_100(self):
        result = self._run(experience_req=0)
        self.assertEqual(result["experience_score"], 100.0)

    # --- education scoring ---

    def test_education_score_100_when_degree_matches(self):
        result = self._run(education_req="Bachelor's")
        self.assertEqual(result["education_score"], 100.0)

    def test_education_score_below_100_when_undereducated(self):
        result = self._run(education_req="Ph.D.")
        self.assertLess(result["education_score"], 100.0)

    def test_no_education_requirement_gives_100(self):
        result = self._run(education_req=None)
        self.assertEqual(result["education_score"], 100.0)

    # --- recommendation label ---

    def test_qualified_label_for_high_scoring_resume(self):
        result = self._run()
        self.assertEqual(result["recommendation_label"], "Qualified")

    def test_not_qualified_label_when_critical_skills_missing(self):
        result = self._run(
            required_skills=["Python", "SQL"],
            critical_skills=["COBOL", "Fortran"],
            requires_all_critical=True,
        )
        self.assertEqual(result["recommendation_label"], "Not Qualified")
        # Skills/exp/edu all 100. Derived from the weights rather than hardcoded
        # so a deliberate reweighting does not silently fail here.
        self.assertEqual(
            result["fit_score"],
            calculate_fit_score(result["skill_score"], result["experience_score"],
                                result["education_score"],
                                result["text_similarity_score"]),
        )
        self.assertEqual(sorted(result["missing_critical_skills"]), sorted(["COBOL", "Fortran"]))
        # Label is forced Not Qualified by the critical-skill rule, but the
        # numeric score is still reported unchanged.
        self.assertIn(f"fit score remains {round(result['fit_score'])}%", result["summary"])

    def test_missing_required_skills_do_not_force_disqualification_without_critical_list(self):
        result = self._run(
            required_skills=["Python", "COBOL"],
            critical_skills=[],
            requires_all_critical=True,
            min_fit_score=50.0,
        )
        self.assertEqual(result["missing_skills"], ["COBOL"])
        self.assertEqual(result["missing_critical_skills"], [])
        self.assertNotEqual(result["recommendation_label"], "Not Qualified")

    def test_critical_skill_enforcement_only_triggers_when_flag_set(self):
        """Missing critical skills should NOT disqualify when flag is False."""
        result = self._run(
            required_skills=["Python", "SQL", "JavaScript"],
            critical_skills=["COBOL"],
            requires_all_critical=False,
            min_fit_score=50.0,
        )
        # COBOL not in resume → skill_score 0 → fit_score low, but label
        # should still be driven by score, not forced.
        self.assertEqual(result["missing_critical_skills"], ["COBOL"])
        self.assertEqual(result["recommendation_label"], "Qualified")

    def test_zero_required_skill_matches_is_not_qualified(self):
        result = self._run(
            required_skills=["COBOL", "Fortran"],
            min_fit_score=50.0,
        )
        self.assertEqual(result["matched_skills"], [])
        self.assertEqual(result["recommendation_label"], "Not Qualified")
        self.assertEqual(
            result["fit_score"],
            calculate_fit_score(result["skill_score"], result["experience_score"],
                                result["education_score"],
                                result["text_similarity_score"]),
        )

    def test_missing_more_than_half_required_skills_routes_to_review(self):
        # The old "< half of required skills -> hard Not Qualified" gate was
        # removed; low skill coverage now simply lowers the weighted fit score,
        # and a borderline result routes to human review rather than auto-failing.
        result = self._run(
            required_skills=["Python", "COBOL", "Fortran", "Haskell"],
            preferred_skills=["Docker"],
            min_fit_score=50.0,
        )
        self.assertEqual(result["matched_skills"], ["Python"])
        self.assertEqual(result["recommendation_label"], "For Review")
        # skills 35 (1 of 4 matched, plus the 10-point preferred bonus), exp/edu 100.
        self.assertEqual(result["skill_score"], 35.0)
        self.assertEqual(
            result["fit_score"],
            calculate_fit_score(result["skill_score"], result["experience_score"],
                                result["education_score"],
                                result["text_similarity_score"]),
        )

    def test_severe_experience_gap_is_not_qualified(self):
        result = self._run(
            experience_req=12,
            min_fit_score=50.0,
        )
        self.assertLess(result["experience_score"], 50.0)
        self.assertEqual(result["recommendation_label"], "Not Qualified")
        # Intent: the weighted score alone would clear the review threshold
        # (fit ~68.95 >= 50), but the experience gate still forces Not Qualified.
        self.assertGreater(result["fit_score"], 50.0)

    def test_partial_experience_gap_routes_candidate_to_review_without_changing_score(self):
        result = self._run(
            experience_req=6,
            min_fit_score=50.0,
        )
        self.assertGreaterEqual(result["experience_score"], 50.0)
        self.assertLess(result["experience_score"], 100.0)
        self.assertEqual(result["recommendation_label"], "For Review")
        self.assertGreater(result["fit_score"], 75.0)

    # --- preferred skills bonus ---

    def test_preferred_skills_increase_skill_score(self):
        base = self._run()
        with_preferred = self._run(preferred_skills=["Docker"])
        # Docker is in the resume, so preferred bonus should be added
        self.assertGreaterEqual(with_preferred["skill_score"], base["skill_score"])

    def test_preferred_skills_capped_at_100(self):
        result = self._run(
            required_skills=["Python"],
            preferred_skills=["SQL", "JavaScript", "React", "Docker"],
        )
        self.assertLessEqual(result["skill_score"], 100.0)

    # --- text similarity ---

    def test_text_similarity_score_between_0_and_100(self):
        result = self._run()
        self.assertGreaterEqual(result["text_similarity_score"], 0.0)
        self.assertLessEqual(result["text_similarity_score"], 100.0)

    # --- contact info ---

    def test_contact_info_email_extracted(self):
        result = self._run()
        self.assertIn("@", result["contact_info"]["email"])

    def test_contact_info_name_extracted(self):
        result = self._run()
        self.assertNotEqual(result["contact_info"]["name"], "")

    # --- summary ---

    def test_summary_is_non_empty_string(self):
        result = self._run()
        self.assertIsInstance(result["summary"], str)
        self.assertGreater(len(result["summary"]), 50)

    def test_decision_explanation_is_transparent_essay(self):
        result = self._run(preferred_skills=["Docker"])
        explanation = result["decision_explanation"]
        self.assertIsInstance(explanation, str)
        self.assertGreater(len(explanation), 100)
        self.assertIn(f"skills {FIT_WEIGHT_PERCENTS[0]}%", explanation)
        self.assertIn(f"experience {FIT_WEIGHT_PERCENTS[1]}%", explanation)
        self.assertIn(f"education {FIT_WEIGHT_PERCENTS[2]}%", explanation)
        self.assertIn(f"text similarity {FIT_WEIGHT_PERCENTS[3]}%", explanation)

    # --- no required skills ---

    def test_empty_required_skills_gives_100_skill_score(self):
        result = self._run(required_skills=[])
        self.assertEqual(result["skill_score"], 100.0)


class TestDecisionConfidence(unittest.TestCase):

    def test_clear_qualified_resume_gets_high_confidence(self):
        level, reason = estimate_decision_confidence(
            recommendation_label="Qualified",
            resume_text=("Maria Santos\nmaria@example.com\nSkills: Python SQL\n"
                         "Experience\nTeacher at City School 2018 - 2024\n"
                         "Education\nBachelor of Education\nState University\n") * 12,
            contact_info={"name": "Maria Santos"},
            extracted_edu=[{"degree": "Bachelor of Education", "institution": "State University"}],
            extracted_exp=[{"job_title": "Teacher", "company": "City School", "years": 6}],
            total_exp_years=6,
            required_skills=["Python", "SQL"],
            matched_skills=["Python", "SQL"],
            missing_skills=[],
            matched_critical_skills=[],
            missing_critical_skills=[],
            skill_score=100,
            exp_score=100,
            edu_score=100,
            fit_score=100,
            experience_req=3,
            education_req="Bachelor's",
        )
        self.assertEqual(level, "High")
        self.assertIn("Confidence support", reason)

    def test_for_review_is_capped_at_medium_confidence(self):
        level, _ = estimate_decision_confidence(
            recommendation_label="For Review",
            resume_text=("Readable resume text with skills and education. " * 50),
            contact_info={"name": "Juan Reyes"},
            extracted_edu=[{"degree": "Bachelor of Science", "institution": "State University"}],
            extracted_exp=[{"job_title": "Analyst", "company": "Acme", "years": 2}],
            total_exp_years=2,
            required_skills=["Excel", "Reports"],
            matched_skills=["Excel"],
            missing_skills=["Reports"],
            matched_critical_skills=[],
            missing_critical_skills=[],
            skill_score=50,
            exp_score=67,
            edu_score=100,
            fit_score=68,
            experience_req=3,
            education_req="Bachelor's",
        )
        self.assertEqual(level, "Medium")

    def test_messy_incomplete_resume_gets_low_confidence(self):
        level, reason = estimate_decision_confidence(
            recommendation_label="Not Qualified",
            resume_text="Java",
            contact_info={"name": "Unknown Candidate"},
            extracted_edu=[],
            extracted_exp=[],
            total_exp_years=0,
            required_skills=["Teaching License", "Classroom Management"],
            matched_skills=[],
            missing_skills=["Teaching License", "Classroom Management"],
            matched_critical_skills=[],
            missing_critical_skills=[],
            skill_score=0,
            exp_score=0,
            edu_score=0,
            fit_score=49,
            experience_req=3,
            education_req="Bachelor's",
        )
        self.assertEqual(level, "Low")
        self.assertIn("Reviewer should verify", reason)

    def test_repeated_text_does_not_create_high_confidence(self):
        level, reason = estimate_decision_confidence(
            recommendation_label="Not Qualified",
            resume_text=("generic resume text " * 100),
            contact_info={"name": "Unknown Candidate"},
            extracted_edu=[], extracted_exp=[], total_exp_years=0,
            required_skills=["Python", "SQL"], matched_skills=[],
            missing_skills=["Python", "SQL"], matched_critical_skills=[],
            missing_critical_skills=[], skill_score=0, exp_score=0,
            edu_score=0, fit_score=20, experience_req=3,
            education_req="Bachelor's",
        )
        self.assertNotEqual(level, "High")
        self.assertIn("sparse or repetitive", reason)

    def test_conflicting_experience_totals_are_flagged(self):
        _, reason = estimate_decision_confidence(
            recommendation_label="Qualified",
            resume_text=("Maria Santos\nExperience\nDeveloper at Acme 2020-2022\n" * 20),
            contact_info={"name": "Maria Santos"}, extracted_edu=[],
            extracted_exp=[{"job_title": "Developer", "company": "Acme", "years": 2}],
            total_exp_years=10, required_skills=["Python"], matched_skills=["Python"],
            missing_skills=[], matched_critical_skills=[], missing_critical_skills=[],
            skill_score=100, exp_score=100, edu_score=100, fit_score=90,
            experience_req=3,
        )
        self.assertIn("conflicts with extracted work-history durations", reason)

    def test_reason_exposes_numeric_evidence_score(self):
        _, reason = estimate_decision_confidence(
            recommendation_label="For Review", resume_text="Short resume",
            contact_info={}, extracted_edu=[], extracted_exp=[], total_exp_years=0,
            required_skills=[], matched_skills=[], missing_skills=[],
            matched_critical_skills=[], missing_critical_skills=[], skill_score=100,
            exp_score=100, edu_score=100, fit_score=60,
        )
        self.assertRegex(reason, r"Evidence confidence score: \d+/100\.")


# ===========================================================================
# File utilities
# ===========================================================================

class TestUniqueUploadFilename(unittest.TestCase):

    def test_returns_unique_filenames(self):
        first = unique_upload_filename("resume.docx")
        second = unique_upload_filename("resume.docx")
        self.assertNotEqual(first, second)

    def test_extension_is_lowercased(self):
        result = unique_upload_filename("../../Resume.DOCX")
        self.assertTrue(result.endswith(".docx"))

    def test_path_traversal_stripped(self):
        result = unique_upload_filename("../../resume.docx")
        self.assertNotIn("..", result)
        self.assertEqual(os.path.basename(result), result)

    def test_only_basename_returned(self):
        result = unique_upload_filename("/some/deep/path/cv.pdf")
        self.assertEqual(os.path.basename(result), result)

    def test_no_original_extension_still_safe(self):
        result = unique_upload_filename("noextension")
        self.assertIsInstance(result, str)
        self.assertGreater(len(result), 0)

    def test_empty_string_input_uses_fallback(self):
        result = unique_upload_filename("")
        self.assertIsInstance(result, str)
        self.assertGreater(len(result), 0)


class TestJobUploadDirectory(unittest.TestCase):

    def test_sorts_by_job_year_month_and_week(self):
        with tempfile.TemporaryDirectory() as upload_root:
            result = job_upload_directory(
                upload_root,
                12,
                "Senior Software Engineer",
                datetime(2026, 8, 9, 10, 30),
            )
            expected = os.path.join(
                os.path.abspath(upload_root),
                "jobs",
                "12-senior-software-engineer",
                "2026",
                "08-august",
                "week-2",
            )
            self.assertEqual(result, expected)

    def test_week_five_is_used_for_end_of_month(self):
        result = job_upload_directory(
            "instance/uploads", 3, "Accountant", datetime(2026, 8, 31)
        )
        self.assertTrue(result.endswith(os.path.join("08-august", "week-5")))

    def test_job_title_is_sanitized_and_cannot_escape_upload_root(self):
        with tempfile.TemporaryDirectory() as upload_root:
            result = job_upload_directory(
                upload_root, 7, "../../HR / Manager", datetime(2026, 1, 1)
            )
            self.assertTrue(is_path_inside_directory(result, upload_root))
            self.assertIn(os.path.join("jobs", "7-hr-manager"), result)

    def test_job_id_separates_duplicate_titles(self):
        uploaded_at = datetime(2026, 8, 9)
        first = job_upload_directory("uploads", 1, "Teacher", uploaded_at)
        second = job_upload_directory("uploads", 2, "Teacher", uploaded_at)
        self.assertNotEqual(first, second)

class TestIsPathInsideDirectory(unittest.TestCase):

    def setUp(self):
        self.upload_dir = os.path.abspath("instance/uploads")

    def test_file_inside_directory_returns_true(self):
        safe_path = os.path.join(self.upload_dir, "resume.docx")
        self.assertTrue(is_path_inside_directory(safe_path, self.upload_dir))

    def test_file_outside_directory_returns_false(self):
        unsafe_path = os.path.abspath("app.py")
        self.assertFalse(is_path_inside_directory(unsafe_path, self.upload_dir))

    def test_path_traversal_attempt_returns_false(self):
        traversal = os.path.join(self.upload_dir, "..", "..", "app.py")
        traversal = os.path.abspath(traversal)
        self.assertFalse(is_path_inside_directory(traversal, self.upload_dir))

    def test_directory_itself_is_inside(self):
        self.assertTrue(is_path_inside_directory(self.upload_dir, self.upload_dir))

    def test_subdirectory_is_inside(self):
        subdir = os.path.join(self.upload_dir, "2024", "january")
        self.assertTrue(is_path_inside_directory(subdir, self.upload_dir))


class TestSafeDeleteUploadedFile(unittest.TestCase):

    def test_deletes_file_inside_allowed_folder(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test_resume.pdf")
            with open(filepath, "w") as f:
                f.write("dummy content")
            result = safe_delete_uploaded_file(filepath, tmpdir)
            self.assertTrue(result)
            self.assertFalse(os.path.exists(filepath))

    def test_refuses_to_delete_file_outside_allowed_folder(self):
        with tempfile.TemporaryDirectory() as allowed_dir:
            with tempfile.NamedTemporaryFile(delete=False) as outside_file:
                outside_path = outside_file.name
            try:
                result = safe_delete_uploaded_file(outside_path, allowed_dir)
                self.assertFalse(result)
                self.assertTrue(os.path.exists(outside_path))
            finally:
                os.unlink(outside_path)

    def test_returns_false_for_nonexistent_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            missing = os.path.join(tmpdir, "ghost.pdf")
            result = safe_delete_uploaded_file(missing, tmpdir)
            self.assertFalse(result)

    def test_returns_false_when_no_filepath_given(self):
        result = safe_delete_uploaded_file(None, "/some/dir")
        self.assertFalse(result)
        result = safe_delete_uploaded_file("", "/some/dir")
        self.assertFalse(result)


# ===========================================================================
# Skill Aliases table sanity checks
# ===========================================================================

class TestSkillAliasesTable(unittest.TestCase):
    """Lightweight smoke-tests that verify the SKILL_ALIASES dict is internally
    consistent and covers the entries the codebase depends on."""

    def test_javascript_has_js_alias(self):
        self.assertIn("js", SKILL_ALIASES.get("javascript", []))

    def test_react_has_reactjs_alias(self):
        self.assertIn("react.js", SKILL_ALIASES.get("react", []))

    def test_html_has_html5_alias(self):
        self.assertIn("html5", SKILL_ALIASES.get("html", []))

    def test_natural_language_processing_has_nlp(self):
        self.assertIn("nlp", SKILL_ALIASES.get("natural language processing", []))

    def test_postgres_and_postgresql_are_cross_referenced(self):
        self.assertIn("postgresql", SKILL_ALIASES.get("postgres", []))
        self.assertIn("postgres", SKILL_ALIASES.get("postgresql", []))


# ===========================================================================
# Entry point
# ===========================================================================

class TestCertificationExtraction(unittest.TestCase):
    def setUp(self):
        from app.services.nlp_pipeline import extract_certifications
        self.extract = extract_certifications

    def test_extracts_board_and_professional_credentials(self):
        text = (
            "CARL ANDREI D. MANALANG, RPm, CPHR\n"
            "Registered Psychometrician  August 2024"
        )
        records = self.extract(text)
        by_name = {record['certification_name']: record for record in records}
        self.assertEqual(by_name['Registered Psychometrician']['date_obtained'], 'August 2024')
        self.assertIn('Certified Professional in Human Resources (CPHR)', by_name)

    def test_normalizes_lpt_and_concatenated_let_passer(self):
        records = self.extract(
            "LicensedProfessionalTeacher\n(LETPasser)\nNovember2021"
        )
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]['certification_name'], 'Licensed Professional Teacher')
        self.assertEqual(records[0]['date_obtained'], 'November2021')

    def test_does_not_invent_credential(self):
        self.assertEqual(
            self.extract('Bachelor of Secondary Education major in English'),
            []
        )


class TestLocalResumeExtractionRegressions(unittest.TestCase):
    def test_malformed_partial_experience_rows_are_discarded(self):
        text = """WORK EXPERIENCES
On-field Staff (Work Immersion)
Central Luzon Drug Rehabilitation Center
Sto. Nino, Pampanga
(A.Y. 2018-2019)
Faculty Member (Teacher)
Holy Child of Mary College Inc.
(S.Y. 2024-2025)
EDUCATION
Bachelor of Education
"""
        records = extract_experience_records(text)

        self.assertEqual(len(records), 2)
        self.assertTrue(all(not record["company"].startswith("(") for record in records))

    def test_compact_education_layout_returns_one_row_per_school(self):
        text = """EDUCATION
Far Eastern University-Manila 2023-Present
Bachelor of Science in Nursing Metro Manila
National University 2021-2023
Senior High School Baliwag, Bulacan
Emigdio A. Bondoc High School 2017-2021
Junior High School San Luis, Pampanga
CERTIFICATIONS, SKILLS & AWARDS
Awards: Graduated Senior High School With High Honors
"""
        records = extract_education(text)

        self.assertEqual(len(records), 3)
        self.assertEqual(records[0]["institution"], "Far Eastern University")
        self.assertEqual(records[1], {
            "degree": "Senior High School",
            "institution": "National University",
        })
        self.assertEqual(records[2], {
            "degree": "Junior High School",
            "institution": "Emigdio A. Bondoc High School",
        })

    def test_company_date_then_role_location_layout(self):
        text = """WORK EXPERIENCE
Far Eastern University - Manila 2024-Present
Student Nurse Manila, Philippines
- Provided quality patient care.
Far Eastern University-Manila 2023-2025
Volunteer Manila, Philippines
- Organized community events.
Far Eastern University-Manila 2023
NSTP Volunteer Bistek Ville 5, Payatas, Quezon City
- Supported outreach activities.
PROJECTS
Research Project
"""
        records = extract_experience_records(text)

        self.assertEqual(len(records), 3)
        self.assertEqual(records[0]["job_title"], "Student Nurse")
        self.assertEqual(records[0]["company"], "Far Eastern University - Manila")
        self.assertEqual(records[0]["location"], "Manila, Philippines")
        self.assertEqual(records[1]["job_title"], "Volunteer")
        self.assertEqual(records[2]["job_title"], "NSTP Volunteer")
        self.assertEqual(records[2]["years"], 1.0)

    def test_deans_lister_is_never_work_experience(self):
        text = """ACHIEVEMENTS
DEAN'S LISTER
College of Education
Academic Years 2020-2024
WORK EXPERIENCES
Faculty Member (Teacher)
Holy Child of Mary College Inc.
(S.Y. 2024 - 2025)
SEMINARS"""
        records = extract_experience_records(text)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]['job_title'], 'Faculty Member (Teacher)')
        self.assertNotIn("LISTER", records[0]['job_title'].upper())

    def test_title_and_company_are_not_reversed(self):
        text = """EXPERIENCE
Infant Jesus Academy - Pampanga  City of San Fernando, Pampanga
Human Resource Management & Development Assistant  July 2024 - Present
SKILLS"""
        record = extract_experience_records(text)[0]
        self.assertEqual(record['job_title'], 'Human Resource Management & Development Assistant')
        self.assertEqual(record['company'], 'Infant Jesus Academy - Pampanga')

    def test_overlapping_roles_are_not_double_counted(self):
        text = """EXPERIENCE
School One
Human Resources Assistant  July 2024 - Present
Freelance
Research Validator  August 2024 - Present
SKILLS"""
        self.assertLessEqual(extract_years_of_experience(text), 2.2)

    def test_teaching_skill_wording_variants_match(self):
        matched, missing = analyze_skills(
            'Designs engaging lesson plans. Implements strong classroom management strategies.',
            ['Lesson Planning', 'Classroom Management'],
        )
        self.assertEqual(matched, ['Lesson Planning', 'Classroom Management'])
        self.assertEqual(missing, [])


class TestExtractionImprovements(unittest.TestCase):
    """Locks in the extraction accuracy fixes (experience headers, proximity
    fallback, and education-domain certifications)."""

    def test_teaching_experience_heading_is_recognized(self):
        text = (
            "Teaching Experience\n"
            "English Teacher, San Jose High School   June 2016 - June 2024\n"
            "Handled lesson planning and classroom management.\n"
            "Education\n"
            "Bachelor of Secondary Education, Holy Angel University 2012 - 2016\n"
        )
        # ~8 years of teaching; the 2012-2016 study period must NOT be counted.
        self.assertEqual(extract_years_of_experience(text), 8.0)

    def test_employment_history_heading_is_recognized(self):
        text = (
            "Employment History\n"
            "Sales Associate at ABC Corp   2019 - 2023\n"
            "Cashier at XYZ Store   2017 - 2019\n"
        )
        self.assertEqual(extract_years_of_experience(text), 6.0)

    def test_proximity_fallback_recovers_job_dates_without_heading(self):
        text = (
            "Juan Dela Cruz\n"
            "Software Engineer, Acme Inc   Jan 2018 - Jan 2023\n"
            "Bachelor of Science in Computer Science, State University   2013 - 2017\n"
        )
        # 5 years from the engineer role; the degree period is excluded.
        self.assertEqual(extract_years_of_experience(text), 5.0)

    def test_school_workplace_is_not_mistaken_for_education(self):
        text = (
            "Ana Cruz\n"
            "Teacher at Manila Science High School  2015 - 2020\n"
            "Studied at University of the Philippines, BS Biology  2010 - 2014\n"
        )
        # Working AT a school is employment; the BS Biology study period is not.
        self.assertEqual(extract_years_of_experience(text), 5.0)

    def test_education_only_resume_yields_zero_experience(self):
        text = (
            "Maria Reyes\n"
            "Bachelor of Elementary Education, City College   2010 - 2014\n"
        )
        self.assertEqual(extract_years_of_experience(text), 0.0)

    def test_new_credentials_are_detected(self):
        for sample, expected in [
            ("Certifications\nTESOL Certificate 2021", "TESOL Certificate"),
            ("National Certificate II in Food Service", "TESDA National Certificate"),
            ("PRC License (Professional Teacher)", "PRC License"),
        ]:
            names = [c["certification_name"] for c in extract_certifications(sample)]
            self.assertIn(expected, names, msg=f"{expected} not found in {names}")

    def test_full_day_dates_and_wrapped_year_keep_work_rows_aligned(self):
        text = """WORK EXPERIENCE
- Remittance clerk (1 year)
Seyer's Freezing Pointe Corporation
Brgy. Lara, CSFP
July 31, 2024 - July 31,
2025
- Auditor (1 year)
Precious Loyal Pet
Brgy. San Agustin, CSFP
Jun 7, 2023 - Jun 22, 2024
- Quality Assurance clerk (7 months)
Superl Philippines Incorporation
Bacolor Pampanga
June 26, 2018 - February 8, 2019
SKILLS
Microsoft Office
"""
        records = extract_experience_records(text)

        self.assertEqual(len(records), 3)
        self.assertEqual(records[0]["job_title"], "Remittance clerk")
        self.assertEqual(records[0]["company"], "Seyer's Freezing Pointe Corporation")
        self.assertEqual(records[0]["location"], "Brgy. Lara, CSFP")
        self.assertEqual(records[0]["years"], 1.0)
        self.assertEqual(records[2]["company"], "Superl Philippines Incorporation")
        self.assertEqual(records[2]["years"], 0.67)

    def test_explicit_school_levels_create_clean_education_rows(self):
        text = """EDUCATIONAL BACKGROUND
Tertiary Level:
Bachelor of Science in Business Administration Major in Marketing
Don Honorio Ventura Technological State University
Secondary Level: San Isidro High School (2014)
Intermediate Education: San Isidro Elementary School (2010)
TRAININGS AND SEMINARS ATTENDED
Pre-employment workshop
"""
        records = extract_education(text)

        self.assertEqual(records, [
            {
                "degree": "Bachelor of Science in Business Administration Major in Marketing",
                "institution": "Don Honorio Ventura Technological State University",
            },
            {"degree": "High School", "institution": "San Isidro High School"},
            {
                "degree": "Elementary Education",
                "institution": "San Isidro Elementary School",
            },
        ])

    def test_repeated_experience_sections_are_combined(self):
        text = """WORK EXPERIENCE
Teacher
Alpha Learning School
January 2020 - January 2022
EDUCATION
Bachelor of Education
State University
TEACHING EXPERIENCE
Tutor
Beta Academy
February 2022 - February 2024
SKILLS
Lesson Planning
"""
        records = extract_experience_records(text)

        self.assertEqual(len(records), 2)
        self.assertEqual(
            {(record["job_title"], record["company"]) for record in records},
            {("Teacher", "Alpha Learning School"), ("Tutor", "Beta Academy")},
        )

    def test_repeated_education_sections_are_combined(self):
        text = """EDUCATION
Bachelor of Elementary Education
First State University
SKILLS
Classroom Management
ACADEMIC QUALIFICATIONS
Master of Arts in Education
Second State University
EXPERIENCE
Teacher
"""
        records = extract_education(text)

        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["institution"], "First State University")
        self.assertEqual(records[1]["institution"], "Second State University")

    def test_empty_explicit_sections_do_not_scan_unrelated_resume_text(self):
        text = """SUMMARY
Teacher role requires a Bachelor's degree and five years of experience.
EDUCATION
SKILLS
Classroom Management
EXPERIENCE
REFERENCES
Available upon request
"""
        self.assertEqual(extract_education(text), [])
        self.assertEqual(extract_experience_records(text), [])


class TestCurrentUploadExtractionRegressions(unittest.TestCase):
    """Generalized regressions derived from the two clean uploaded resumes."""

    def test_level_labelled_education_keeps_every_school(self):
        text = """EDUCATION
Primary
Salacot Elementary School
2007-2008
Bancal Elementary School
2008-2013
Secondary
Sto. Tomas High School
2013-2017
Mary The Queen College
2017-2019
Tertiary
Guagua National Colleges Inc.
2019-2023
WORK EXPERIENCE
Student Teacher
"""

        records = extract_education(text)

        self.assertEqual(
            [(record["degree"], record["institution"]) for record in records],
            [
                ("Elementary School", "Salacot Elementary School"),
                ("Elementary School", "Bancal Elementary School"),
                ("High School", "Sto. Tomas High School"),
                ("High School", "Mary The Queen College"),
                ("Tertiary Education", "Guagua National Colleges Inc"),
            ],
        )

    def test_wrapped_certificate_ending_in_preposition_is_one_record(self):
        records = extract_certifications("""CERTIFICATES
Member of the Technical Working Group in
the 2023 CSSPC
LANGUAGES
English
""")

        self.assertEqual(len(records), 1)
        self.assertEqual(
            records[0]["certification_name"],
            "Member of the Technical Working Group in the CSSPC",
        )
        self.assertEqual(records[0]["date_obtained"], "2023")

    def test_scholastic_and_working_experience_boundaries_prevent_false_rows(self):
        text = """SCHOLASTIC RECORDS:
Master's Degree: Master of Arts in Education Major in Filipino
Don Honorio Ventura State University
2021 - 2022
Tertiary: Bachelor in Secondary Education
Major in Filipino
Don Honorio Ventura Technological State University
2014 - 2018
Secondary: Lubao National High School
2009 - 2013
Elementary: Sto. Tomas Elementary School
2003 - 2009
SEMINARS and TRAININGS ATTENDED:
Participant, 1st National Research Conference for Teacher Education
Don Honorio Ventura State University
December 14, 2021
WORKING EXPERIENCE:
College Instructor I
BSEd Filipino Coordinator
Newsletter Adviser (LUALU)
September 1, 2022 - PRESENT
Don Honorio Ventura State University
Senior High School Teacher
June 2019 - May 2022
GUAGUA NATIONAL COLLEGES Inc.
Volunteer Teacher
Summer Pre-Kindergarten Program (SPKP)
April - May 2019
New Era University
Quezon City, Philippines
Senior High School Teacher
June 2018 - March 2019
AMA Computer Learning Center Inc.
PERSONAL DATA:
Age: 26
"""

        education = extract_education(text)
        experience = extract_experience_records(text)

        self.assertEqual(len(education), 4)
        self.assertEqual(len(experience), 4)
        self.assertEqual(
            experience[0]["job_title"],
            "College Instructor I / BSEd Filipino Coordinator / Newsletter Adviser (LUALU)",
        )
        self.assertEqual(experience[0]["company"], "Don Honorio Ventura State University")
        self.assertIn(
            {
                "job_title": "Volunteer Teacher",
                "company": "New Era University",
                "location": "Quezon City, Philippines",
                "years": 0.08,
            },
            experience,
        )
        false_companies = {
            "BSEd Filipino Coordinator",
            "Cabambangan Villa de Bacolor",
            "Summer Pre-Kindergarten Program (SPKP)",
        }
        self.assertTrue(false_companies.isdisjoint({record["company"] for record in experience}))

    def test_header_name_is_not_reconstructed_from_na_project_text(self):
        text = """HERO D. PARK
Angeles City, Pampanga | hdpark09@gmail.com | (+63) 09455261300
PROJECTS
Attendance Application
N/A
"""
        contact = extract_contact_info(text)

        self.assertEqual(contact["name"], "Hero D. Park")
        self.assertEqual(contact["phone"], "09455261300")

    def test_location_and_reference_phone_are_not_applicant_contact_data(self):
        text = """GENE ELPIE L. LANDOY
Graphics Designer | UI/UX Designer
portfolio.example | Boac Marinduque | g.landoyelpie@gmail.com
CHARACTER REFERENCE
Doreena Joy Borja, LPT
State University
09171234567
"""
        contact = extract_contact_info(text)

        self.assertEqual(contact["name"], "Gene Elpie L. Landoy")
        self.assertEqual(contact["phone"], "Unknown Phone")

    def test_institution_does_not_absorb_trailing_city_and_country(self):
        records = extract_education("""EDUCATION
Holy Angel University Angeles City, Philippines
Bachelor of Science in Computer Science - Dean's Lister
PROJECTS
Attendance Application
""")

        self.assertEqual(records[0]["institution"], "Holy Angel University")

    def test_role_date_then_employer_layout_returns_complete_history(self):
        text = """PROFESSIONAL EXPERIENCE
Backend Developer and Product Owner (Intern) | Jan 2026 - Apr 2026
Department of Science and Technology - Regional Developers
Built backend systems and coordinated delivery.
Developer and UI/UX Designer | Jul 2024 - Jan 2026
College of Information and Computing Sciences - DevTeam
Improved technical proficiency and real-world project experience.
Infocus Correspondent | Jul 2025 - Jan 2026
Sentro Publication
Produced campus news.
Associate Editor-in-Chief | Jul 2025 - Apr 2026
Infocus Publication and Broadcasting
Improved workflow efficiency and output delivery.
Editor-in-Chief | Jul 2024 - Jul 2025
Infocus Publication and Broadcasting
Enhanced audience reach and visual storytelling.
Head Layout Artist | Jul 2022 - Jul 2024
Infocus Publication and Broadcasting
Designed publication layouts.
CERTIFICATES
Course Certificate
"""
        records = extract_experience_records(text)

        self.assertEqual(len(records), 6)
        self.assertEqual(records[0]["job_title"], "Backend Developer and Product Owner (Intern)")
        self.assertEqual(records[-1]["job_title"], "Head Layout Artist")
        self.assertNotIn("proficiency and real-world project experience", {
            record["company"] for record in records
        })

    def test_wrapped_pipe_certifications_are_rejoined(self):
        text = """CERTIFICATIONS
Responsive Web Design Certification | CCNA: Introduction to Networks | Data Analytics Essentials | AI
Fundamentals with IBM SkillsBuild | Cyber Threat Management | JavaScript Essentials 1 |
CompTIA IT Fundamentals (ITF+)
"""
        names = [record["certification_name"] for record in extract_certifications(text)]

        self.assertEqual(len(names), 7)
        self.assertIn("AI Fundamentals with IBM SkillsBuild", names)
        self.assertIn("CompTIA IT Fundamentals (ITF+)", names)

    def test_certificate_section_stops_before_projects_and_references(self):
        text = """CERTIFICATES
● Civil Service Eligibility
● Web Development Workshop Certification (Developer Club, Coding
Bootcamp, 2023)
RELEVANT COURSEWORK AND PROJECTS
AI Resume Parser (Python, NLP)
CHARACTER REFERENCE
Alex Reyes, LPT
09171234567
"""
        records = extract_certifications(text)
        names = [record["certification_name"] for record in records]

        self.assertIn("Civil Service Eligibility", names)
        self.assertIn(
            "Web Development Workshop Certification (Developer Club, Coding Bootcamp)",
            names,
        )
        self.assertNotIn("AI Resume Parser (Python, NLP)", names)
        self.assertNotIn("Licensed Professional Teacher", names)

    def test_wrapped_skill_categories_preserve_compound_items(self):
        skills = extract_resume_skills("""TECHNICAL SKILLS
Frameworks & Libraries: Flutter (Mobile & Web), FastAPI, NumPy,
Scikit-Learn
Core Concepts: Machine Learning, Computer Vision (CNNs), API
Integration
CERTIFICATIONS
Responsive Web Design
""")

        self.assertIn("Scikit-Learn", skills)
        self.assertIn("API Integration", skills)
        self.assertNotIn("Frameworks & Libraries", skills)
        self.assertNotIn("API", skills)
        self.assertNotIn("Integration", skills)


if __name__ == "__main__":
    unittest.main(verbosity=2)
